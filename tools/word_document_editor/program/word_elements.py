from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from word_formula import latex_to_word_omml, preprocess_latex
from word_table_width import (
    CELL_HORIZONTAL_PADDING_POINTS,
    FontTextMeasurer,
    apply_column_widths,
    calculate_column_widths,
    document_available_width_points,
    set_cell_margins,
    set_repeat_header,
)


DEFAULT_THEME = {
    "font_family": "Times New Roman",
    "east_asia_font": "微软雅黑",
    "complex_script_font": "Times New Roman",
    "heading_color": "000000",
    "text_color": "000000",
    "muted_color": "64748B",
    "accent_color": "00A6A6",
    "table_header_color": "F2F2F2",
    "table_header_text_color": "000000",
}
BODY_FONT_SIZE = 12.0
TABLE_FONT_SIZE = 10.5
BODY_SPACE_AFTER = 6.0
BODY_LINE_SPACING = 1.15
BODY_FIRST_LINE_INDENT = 0.28
HEADING_FONT_SIZES = {1: 18.0, 2: 15.0, 3: 13.0}
FORMULA_XSL_PATH = Path(__file__).resolve().parents[1] / "assets" / "MML2OMML.XSL"


def merged_theme(value: Any) -> dict[str, Any]:
    theme = dict(DEFAULT_THEME)
    if isinstance(value, dict):
        theme.update({key: item for key, item in value.items() if item is not None})
    normalize_theme_fonts(theme)
    return theme


def normalize_theme_fonts(theme: dict[str, Any]) -> None:
    font_family = str(theme.get("font_family") or DEFAULT_THEME["font_family"])
    east_asia_font = (
        theme.get("east_asia_font")
        or theme.get("eastAsia_font")
        or theme.get("cjk_font")
        or font_family
    )
    complex_script_font = theme.get("complex_script_font") or theme.get("cs_font") or font_family
    theme["font_family"] = font_family
    theme["east_asia_font"] = str(east_asia_font)
    theme["complex_script_font"] = str(complex_script_font)


def set_header_footer(doc: Document, *, header: str | None = None, footer: str | None = None, theme: dict[str, Any]) -> None:
    for section in doc.sections:
        if header is not None:
            paragraph = section.header.paragraphs[0]
            paragraph.text = header
            apply_paragraph_style(paragraph, {"font_size": 9, "color": theme["muted_color"]}, theme)
        if footer is not None:
            paragraph = section.footer.paragraphs[0]
            paragraph.text = footer
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            apply_paragraph_style(paragraph, {"font_size": 9, "color": theme["muted_color"]}, theme)


def add_elements(
    doc: Document,
    elements: list[Any],
    theme: dict[str, Any],
    root: Path,
    warnings: list[str] | None = None,
) -> dict[str, int]:
    stats = {
        "headings": 0,
        "paragraphs": 0,
        "lists": 0,
        "tables": 0,
        "images": 0,
        "equations": 0,
        "equation_fallbacks": 0,
        "page_breaks": 0,
    }
    for element in elements:
        if not isinstance(element, dict):
            continue
        element_type = str(element.get("type") or "").lower()
        if element_type == "heading":
            add_heading(doc, element, theme)
            stats["headings"] += 1
        elif element_type == "paragraph":
            add_paragraph(doc, element, theme)
            stats["paragraphs"] += 1
        elif element_type == "bullets":
            add_list(doc, element, theme, numbered=False)
            stats["lists"] += 1
        elif element_type == "numbered":
            add_list(doc, element, theme, numbered=True)
            stats["lists"] += 1
        elif element_type == "table":
            add_table(doc, element, theme)
            stats["tables"] += 1
        elif element_type == "image":
            add_image(doc, element, root)
            stats["images"] += 1
        elif element_type == "equation":
            used_fallback = add_equation(doc, element, theme, warnings=warnings)
            stats["equations"] += 1
            stats["equation_fallbacks"] += 1 if used_fallback else 0
        elif element_type == "page_break":
            doc.add_page_break()
            stats["page_breaks"] += 1
        else:
            raise ValueError(f"不支持的 Word 元素类型：{element_type}")
    return stats


def add_heading(doc: Document, spec: dict[str, Any], theme: dict[str, Any]) -> Any:
    level = spec.get("level")
    if not isinstance(level, int):
        level = 1
    level = max(1, min(level, 9))
    paragraph = doc.add_heading(str(spec.get("text") or ""), level=level)
    style = {
        "font_size": HEADING_FONT_SIZES.get(level, BODY_FONT_SIZE),
        "bold": True,
        "color": theme["heading_color"],
        "space_before": 12 if level <= 2 else 8,
        "space_after": 6 if level <= 3 else 4,
        "line_spacing": BODY_LINE_SPACING,
        "keep_with_next": True,
    }
    if isinstance(spec.get("style"), dict):
        style.update(spec["style"])
    apply_paragraph_style(paragraph, style, theme)
    return paragraph


def add_paragraph(doc: Document, spec: dict[str, Any], theme: dict[str, Any]) -> Any:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(str(spec.get("text") or ""))
    style = {
        "font_size": BODY_FONT_SIZE,
        "color": theme["text_color"],
        "space_before": 0,
        "space_after": BODY_SPACE_AFTER,
        "line_spacing": BODY_LINE_SPACING,
        "first_line_indent": BODY_FIRST_LINE_INDENT,
    }
    if isinstance(spec.get("style"), dict):
        style.update(spec["style"])
    apply_font(run.font, style, theme)
    apply_paragraph_format(paragraph, style)
    return paragraph


def add_list(doc: Document, spec: dict[str, Any], theme: dict[str, Any], *, numbered: bool) -> None:
    items = spec.get("items") or []
    if not isinstance(items, list):
        raise ValueError("列表元素 items 必须是数组。")
    base_style = {
        "font_size": BODY_FONT_SIZE,
        "color": theme["text_color"],
        "space_before": 0,
        "space_after": BODY_SPACE_AFTER,
        "line_spacing": BODY_LINE_SPACING,
    }
    if isinstance(spec.get("style"), dict):
        base_style.update(spec["style"])
    for item in items:
        if isinstance(item, dict):
            text = str(item.get("text") or "")
            level = int(item.get("level") or 0)
        else:
            text = str(item)
            level = 0
        style_name = "List Number" if numbered else "List Bullet"
        if level > 0:
            style_name += f" {min(level + 1, 3)}"
        paragraph = doc.add_paragraph(style=style_name)
        run = paragraph.add_run(text)
        apply_font(run.font, base_style, theme)
        apply_paragraph_format(paragraph, base_style)


def add_table(doc: Document, spec: dict[str, Any], theme: dict[str, Any]) -> Any:
    rows = spec.get("rows")
    if not isinstance(rows, list) or not rows:
        raise ValueError("table 元素必须提供非空 rows。")
    col_count = max(len(row) for row in rows if isinstance(row, list))
    if col_count <= 0:
        raise ValueError("table.rows 不能为空。")
    table = doc.add_table(rows=len(rows), cols=col_count)
    style = spec.get("style") if isinstance(spec.get("style"), dict) else {}
    table.style = str(style.get("table_style") or "Table Grid")
    table.alignment = table_alignment(style.get("align") or "left")
    normalized_rows = [
        ["" if index >= len(row) or row[index] is None else str(row[index]) for index in range(col_count)]
        if isinstance(row, list)
        else [""] * col_count
        for row in rows
    ]
    explicit_widths = style.get("column_widths")
    if isinstance(explicit_widths, list) and len(explicit_widths) == col_count:
        try:
            column_widths = [max(0.0, float(value)) for value in explicit_widths]
        except (TypeError, ValueError):
            column_widths = []
    else:
        column_widths = []
    if not column_widths or sum(column_widths) <= 0:
        font_name = str(style.get("font_family") or theme["font_family"])
        font_size = float(style.get("font_size") or TABLE_FONT_SIZE)
        with FontTextMeasurer(font_name=font_name, size_points=font_size) as measurer:
            column_widths = calculate_column_widths(
                normalized_rows[0],
                normalized_rows[1:],
                available_width_points=document_available_width_points(doc),
                cell_padding_points=CELL_HORIZONTAL_PADDING_POINTS,
                measurer=measurer,
            )
    apply_column_widths(table, column_widths, doc)
    set_repeat_header(table.rows[0])
    for row_index, row in enumerate(rows):
        if not isinstance(row, list):
            continue
        for col_index in range(col_count):
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            cell.text = normalized_rows[row_index][col_index]
            cell_style = {
                "font_size": TABLE_FONT_SIZE,
                "color": theme["text_color"],
                "space_before": 0,
                "space_after": 0,
                "line_spacing": 1.0,
            }
            cell_style.update(style)
            if row_index == 0:
                shade_cell(cell, str(style.get("header_fill") or theme["table_header_color"]))
                cell_style.update({
                    "bold": True,
                    "color": style.get("header_color") or theme["table_header_text_color"],
                })
                apply_cell_text_style(cell, cell_style, theme)
            else:
                apply_cell_text_style(cell, cell_style, theme)
    return table


def add_image(doc: Document, spec: dict[str, Any], root: Path) -> None:
    raw_path = spec.get("image_path")
    if not isinstance(raw_path, str) or not raw_path.strip():
        raise ValueError("image 元素缺少 image_path。")
    image_path = Path(raw_path).expanduser()
    if not image_path.is_absolute():
        image_path = root / image_path
    image_path = image_path.resolve(strict=False)
    try:
        image_path.relative_to(root)
    except ValueError as exc:
        raise ValueError(f"图片路径不在工作区内：{image_path}") from exc
    if not image_path.is_file():
        raise ValueError(f"图片文件不存在：{image_path}")
    width = spec.get("width")
    if isinstance(width, (int, float)):
        doc.add_picture(str(image_path), width=Inches(width))
    else:
        doc.add_picture(str(image_path))


def add_equation(doc: Document, spec: dict[str, Any], theme: dict[str, Any], warnings: list[str] | None = None) -> bool:
    latex = spec.get("latex")
    if not isinstance(latex, str) or not latex.strip():
        latex = spec.get("text")
    if not isinstance(latex, str) or not latex.strip():
        raise ValueError("equation 元素必须提供非空 latex。")

    style = spec.get("style") if isinstance(spec.get("style"), dict) else {}
    display = spec.get("display")
    display = display if isinstance(display, bool) else True
    paragraph = doc.add_paragraph()
    paragraph_style = dict(style)
    paragraph_style.setdefault("align", "center" if display else "left")
    paragraph_style.setdefault("space_before", 0)
    paragraph_style.setdefault("space_after", BODY_SPACE_AFTER)
    paragraph_style.setdefault("line_spacing", BODY_LINE_SPACING)
    apply_paragraph_format(paragraph, paragraph_style)
    return add_formula_to_paragraph(paragraph, latex, theme, warnings=warnings)


def add_formula_to_paragraph(paragraph: Any, latex: str, theme: dict[str, Any], warnings: list[str] | None = None) -> bool:
    normalized = preprocess_latex(latex)
    if not normalized:
        return False
    if len(normalized) <= 3 and normalized.isalnum():
        run = paragraph.add_run(normalized)
        run.italic = True
        set_font_family(run.font, "Times New Roman", east_asia_font=theme["east_asia_font"])
        run.font.size = Pt(BODY_FONT_SIZE)
        return False

    omml, error = latex_to_word_omml(latex, xsl_path=FORMULA_XSL_PATH, warnings=warnings)
    if omml is not None:
        paragraph._p.append(omml)
        return False

    if warnings is not None:
        warnings.append(f"公式降级为文本：{normalized}，原因：{error}")
    run = paragraph.add_run(f"${normalized}$")
    set_font_family(run.font, "Consolas", east_asia_font=theme["east_asia_font"])
    run.font.size = Pt(10)
    run.font.color.rgb = parse_color(theme["muted_color"])
    return True


def apply_paragraph_style(paragraph: Any, style: dict[str, Any], theme: dict[str, Any]) -> None:
    for run in paragraph.runs:
        apply_font(run.font, style, theme)
    apply_paragraph_format(paragraph, style)


def apply_cell_text_style(cell: Any, style: dict[str, Any], theme: dict[str, Any]) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            apply_font(run.font, style, theme)
        apply_paragraph_format(paragraph, style)


def apply_font(font: Any, style: dict[str, Any], theme: dict[str, Any]) -> None:
    font_family = str(style.get("font_family") or theme["font_family"])
    east_asia_font = str(
        style.get("east_asia_font")
        or style.get("eastAsia_font")
        or style.get("cjk_font")
        or theme["east_asia_font"]
        or font_family
    )
    complex_script_font = str(
        style.get("complex_script_font")
        or style.get("cs_font")
        or theme["complex_script_font"]
        or font_family
    )
    set_font_family(
        font,
        font_family,
        east_asia_font=east_asia_font,
        complex_script_font=complex_script_font,
    )
    font.size = Pt(float(style.get("font_size") or BODY_FONT_SIZE))
    if "bold" in style:
        font.bold = bool(style["bold"])
    if "italic" in style:
        font.italic = bool(style["italic"])
    font.color.rgb = parse_color(style.get("color") or theme["text_color"])


def set_font_family(
    font: Any,
    font_family: str,
    *,
    east_asia_font: str | None = None,
    complex_script_font: str | None = None,
) -> None:
    font.name = font_family
    font_element = font_properties_element(font)
    if font_element is None:
        return
    r_fonts = font_element.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        font_element.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), font_family)
    r_fonts.set(qn("w:hAnsi"), font_family)
    r_fonts.set(qn("w:eastAsia"), east_asia_font or font_family)
    r_fonts.set(qn("w:cs"), complex_script_font or font_family)


def font_properties_element(font: Any) -> Any:
    font_element = getattr(font, "_element", None)
    if font_element is None:
        return None
    if hasattr(font_element, "rFonts"):
        return font_element
    if hasattr(font_element, "get_or_add_rPr"):
        return font_element.get_or_add_rPr()
    return None


def apply_paragraph_format(paragraph: Any, style: dict[str, Any]) -> None:
    paragraph.alignment = paragraph_alignment(style.get("align"))
    if isinstance(style.get("space_after"), (int, float)):
        paragraph.paragraph_format.space_after = Pt(style["space_after"])
    if isinstance(style.get("space_before"), (int, float)):
        paragraph.paragraph_format.space_before = Pt(style["space_before"])
    if isinstance(style.get("line_spacing"), (int, float)):
        paragraph.paragraph_format.line_spacing = style["line_spacing"]
    if isinstance(style.get("first_line_indent"), (int, float)):
        paragraph.paragraph_format.first_line_indent = Inches(style["first_line_indent"])
    if "keep_with_next" in style:
        paragraph.paragraph_format.keep_with_next = bool(style["keep_with_next"])


def paragraph_alignment(value: Any) -> Any:
    mapping = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    return mapping.get(str(value or "left").lower(), WD_ALIGN_PARAGRAPH.LEFT)


def table_alignment(value: Any) -> Any:
    mapping = {
        "left": WD_TABLE_ALIGNMENT.LEFT,
        "center": WD_TABLE_ALIGNMENT.CENTER,
        "right": WD_TABLE_ALIGNMENT.RIGHT,
    }
    return mapping.get(str(value or "left").lower(), WD_TABLE_ALIGNMENT.LEFT)


def shade_cell(cell: Any, color: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), normalize_color(color))
    cell._tc.get_or_add_tcPr().append(shading)


def parse_color(value: Any) -> RGBColor:
    raw = normalize_color(value)
    return RGBColor(int(raw[0:2], 16), int(raw[2:4], 16), int(raw[4:6], 16))


def normalize_color(value: Any) -> str:
    raw = str(value or "000000").strip().lstrip("#")
    if len(raw) == 8:
        raw = raw[-6:]
    if len(raw) != 6 or any(char not in "0123456789abcdefABCDEF" for char in raw):
        raw = "000000"
    return raw.upper()
