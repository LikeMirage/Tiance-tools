from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx.enum.text import WD_BREAK
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from word_elements import (
    BODY_FIRST_LINE_INDENT,
    BODY_FONT_SIZE,
    BODY_LINE_SPACING,
    BODY_SPACE_AFTER,
    HEADING_FONT_SIZES,
    TABLE_FONT_SIZE,
    add_formula_to_paragraph,
    add_image,
    add_table,
    apply_font,
    apply_paragraph_format,
)


TABLE_SEPARATOR_RE = re.compile(r"^:?-{2,}:?$")
LIST_RE = re.compile(r"^(\s*)([-*+]|\d+[.)])\s+(.+)$")
IMAGE_RE = re.compile(r"^!\[([^]]*)\]\((.+)\)$")
LINK_RE = re.compile(r"^\[([^]]+)\]\(([^)]+)\)")
CONTROL_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")


def append_markdown_fragment(
    doc: Any,
    markdown: str,
    theme: dict[str, Any],
    root: Path,
    warnings: list[str],
) -> dict[str, int]:
    if not isinstance(markdown, str) or not markdown.strip():
        raise ValueError("Markdown 写入内容不能为空。")
    cleaned, replacements = CONTROL_RE.subn("�", markdown.replace("\r\n", "\n").replace("\r", "\n"))
    if replacements:
        warnings.append(f"Markdown 中 {replacements} 个无效控制字符已替换。")
    lines = cleaned.split("\n")
    stats = empty_stats()
    index = 0
    while index < len(lines):
        stripped = lines[index].strip().lstrip("\ufeff")
        if not stripped:
            index += 1
            continue
        if stripped in {"<!-- pagebreak -->", "\\pagebreak"}:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            stats["page_breaks"] += 1
            index += 1
            continue
        if stripped.startswith("$$") or stripped.startswith(r"\["):
            latex, index = collect_block_formula(lines, index)
            paragraph = doc.add_paragraph()
            apply_paragraph_format(paragraph, {"align": "center"})
            if add_formula_to_paragraph(paragraph, latex, theme, warnings=warnings):
                stats["equation_fallbacks"] += 1
            stats["equations"] += 1
            continue
        heading = re.match(r"^(#{1,9})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            paragraph = doc.add_heading(level=level)
            style = {
                "font_size": HEADING_FONT_SIZES.get(level, BODY_FONT_SIZE),
                "bold": True,
                "color": theme["heading_color"],
                "space_before": 12 if level <= 2 else 8,
                "space_after": 6 if level <= 3 else 4,
                "line_spacing": BODY_LINE_SPACING,
                "keep_with_next": True,
            }
            apply_paragraph_format(paragraph, style)
            write_inline(paragraph, heading.group(2), style, theme, root, warnings, stats)
            stats["headings"] += 1
            index += 1
            continue
        if is_table_start(lines, index):
            index = add_markdown_table(doc, lines, index, theme, root, warnings, stats)
            continue
        list_match = LIST_RE.match(lines[index])
        if list_match:
            index = add_markdown_list(doc, lines, index, theme, root, warnings, stats)
            continue
        if stripped.startswith(">"):
            index = add_blockquote(doc, lines, index, theme, root, warnings, stats)
            continue
        if stripped.startswith(("```", "~~~")):
            index = add_code_block(doc, lines, index, theme, stats)
            continue
        image = IMAGE_RE.match(stripped)
        if image:
            add_image(doc, {"type": "image", "image_path": parse_destination(image.group(2))}, root)
            stats["images"] += 1
            index += 1
            continue
        if stripped in {"---", "***", "___"}:
            add_horizontal_rule(doc)
            stats["paragraphs"] += 1
            index += 1
            continue
        paragraph_lines = [stripped]
        index += 1
        while index < len(lines) and lines[index].strip() and not starts_block(lines, index):
            paragraph_lines.append(lines[index].strip())
            index += 1
        paragraph = doc.add_paragraph()
        style = {
            "font_size": BODY_FONT_SIZE,
            "color": theme["text_color"],
            "space_before": 0,
            "space_after": BODY_SPACE_AFTER,
            "line_spacing": BODY_LINE_SPACING,
            "first_line_indent": BODY_FIRST_LINE_INDENT,
        }
        apply_paragraph_format(paragraph, style)
        write_inline(paragraph, " ".join(paragraph_lines), style, theme, root, warnings, stats)
        stats["paragraphs"] += 1
    return stats


def write_inline(
    paragraph: Any,
    text: str,
    base_style: dict[str, Any],
    theme: dict[str, Any],
    root: Path,
    warnings: list[str],
    stats: dict[str, int],
) -> None:
    position = 0
    plain: list[str] = []

    def flush() -> None:
        if not plain:
            return
        add_styled_run(paragraph, "".join(plain), base_style, theme)
        plain.clear()

    while position < len(text):
        if text[position] == "\\" and position + 1 < len(text) and text[position + 1] in r"\`*_{}[]()#+-.!$|>":
            plain.append(text[position + 1])
            position += 2
            continue
        token = inline_token(text, position)
        if token is None:
            plain.append(text[position])
            position += 1
            continue
        flush()
        end, kind, value, extra = token
        if kind == "math":
            if add_formula_to_paragraph(paragraph, value, theme, warnings=warnings):
                stats["equation_fallbacks"] += 1
            stats["equations"] += 1
        elif kind == "link":
            add_hyperlink(paragraph, value, extra, base_style, theme)
        else:
            style = dict(base_style)
            if kind == "bold":
                style["bold"] = True
            elif kind == "italic":
                style["italic"] = True
            elif kind == "code":
                style.update({"font_family": "Consolas", "east_asia_font": "Consolas"})
            run = add_styled_run(paragraph, value, style, theme)
            if kind == "strike":
                run.font.strike = True
        position = end
    flush()


def inline_token(text: str, position: int) -> tuple[int, str, str, str] | None:
    if text.startswith("**", position) or text.startswith("__", position):
        delimiter = text[position : position + 2]
        end = find_unescaped(text, delimiter, position + 2)
        if end >= 0:
            return end + 2, "bold", text[position + 2 : end], ""
    if text.startswith("~~", position):
        end = find_unescaped(text, "~~", position + 2)
        if end >= 0:
            return end + 2, "strike", text[position + 2 : end], ""
    if text[position] == "`":
        end = find_unescaped(text, "`", position + 1)
        if end >= 0:
            return end + 1, "code", text[position + 1 : end], ""
    if text.startswith(r"\(", position):
        end = find_unescaped(text, r"\)", position + 2)
        if end >= 0:
            return end + 2, "math", text[position + 2 : end], ""
    if text[position] == "$" and not text.startswith("$$", position):
        end = find_unescaped(text, "$", position + 1)
        if end >= 0 and text[position + 1 : end].strip() == text[position + 1 : end]:
            return end + 1, "math", text[position + 1 : end], ""
    if text[position] == "[":
        match = LINK_RE.match(text[position:])
        if match:
            return position + match.end(), "link", match.group(1), parse_destination(match.group(2))
    if text[position] in {"*", "_"}:
        delimiter = text[position]
        end = find_unescaped(text, delimiter, position + 1)
        if end > position + 1:
            return end + 1, "italic", text[position + 1 : end], ""
    return None


def add_markdown_table(
    doc: Any,
    lines: list[str],
    start: int,
    theme: dict[str, Any],
    root: Path,
    warnings: list[str],
    stats: dict[str, int],
) -> int:
    rows = [parse_table_row(lines[start])]
    index = start + 2
    while index < len(lines) and "|" in lines[index] and lines[index].strip():
        rows.append(parse_table_row(lines[index]))
        index += 1
    width = max(len(row) for row in rows)
    normalized = [(row + [""] * width)[:width] for row in rows]
    table = add_table(doc, {"type": "table", "rows": normalized}, theme)
    for row_index, row in enumerate(normalized):
        for column_index, value in enumerate(row):
            cell = table.cell(row_index, column_index)
            paragraph = cell.paragraphs[0]
            paragraph.clear()
            style = {
                "font_size": TABLE_FONT_SIZE,
                "bold": row_index == 0,
                "color": theme["table_header_text_color"] if row_index == 0 else theme["text_color"],
                "space_before": 0,
                "space_after": 0,
                "line_spacing": 1.0,
            }
            apply_paragraph_format(paragraph, style)
            write_inline(paragraph, value, style, theme, root, warnings, stats)
    stats["tables"] += 1
    return index


def add_markdown_list(
    doc: Any,
    lines: list[str],
    start: int,
    theme: dict[str, Any],
    root: Path,
    warnings: list[str],
    stats: dict[str, int],
) -> int:
    index = start
    while index < len(lines):
        match = LIST_RE.match(lines[index])
        if not match:
            break
        marker = match.group(2)
        ordered = marker[0].isdigit()
        level = min(len(match.group(1).replace("\t", "    ")) // 2, 2)
        style_name = "List Number" if ordered else "List Bullet"
        if level:
            style_name += f" {level + 1}"
        paragraph = doc.add_paragraph(style=style_name)
        write_inline(
            paragraph,
            match.group(3),
            {
                "font_size": BODY_FONT_SIZE,
                "color": theme["text_color"],
                "space_before": 0,
                "space_after": BODY_SPACE_AFTER,
                "line_spacing": BODY_LINE_SPACING,
            },
            theme,
            root,
            warnings,
            stats,
        )
        stats["list_items"] += 1
        index += 1
    return index


def add_blockquote(
    doc: Any,
    lines: list[str],
    start: int,
    theme: dict[str, Any],
    root: Path,
    warnings: list[str],
    stats: dict[str, int],
) -> int:
    index = start
    while index < len(lines) and lines[index].strip().startswith(">"):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Pt(18)
        text = re.sub(r"^\s*>\s?", "", lines[index])
        write_inline(
            paragraph,
            text,
            {"font_size": BODY_FONT_SIZE, "italic": True, "color": theme["muted_color"]},
            theme,
            root,
            warnings,
            stats,
        )
        stats["quotes"] += 1
        index += 1
    return index


def add_code_block(doc: Any, lines: list[str], start: int, theme: dict[str, Any], stats: dict[str, int]) -> int:
    marker = lines[start].strip()[:3]
    index = start + 1
    code_lines: list[str] = []
    while index < len(lines) and not lines[index].strip().startswith(marker):
        code_lines.append(lines[index])
        index += 1
    if index < len(lines):
        index += 1
    paragraph = doc.add_paragraph()
    for line_index, line in enumerate(code_lines or [""]):
        if line_index:
            paragraph.add_run().add_break()
        run = add_styled_run(
            paragraph,
            line or " ",
            {"font_family": "Consolas", "east_asia_font": "Consolas", "font_size": 9, "color": theme["text_color"]},
            theme,
        )
        run.font.name = "Consolas"
    stats["code_blocks"] += 1
    return index


def collect_block_formula(lines: list[str], start: int) -> tuple[str, int]:
    stripped = lines[start].strip()
    opener, closer = ("$$", "$$") if stripped.startswith("$$") else (r"\[", r"\]")
    tail = stripped[len(opener) :]
    if tail.endswith(closer) and len(tail) > len(closer):
        return tail[: -len(closer)].strip(), start + 1
    parts = [tail] if tail else []
    index = start + 1
    while index < len(lines):
        current = lines[index].strip()
        if current.endswith(closer):
            if current != closer:
                parts.append(current[: -len(closer)])
            latex = "\n".join(parts).strip()
            if not latex:
                raise ValueError("Markdown 块公式不能为空。")
            return latex, index + 1
        parts.append(lines[index])
        index += 1
    raise ValueError("Markdown 块公式定界符未闭合。")


def is_table_start(lines: list[str], index: int) -> bool:
    if index + 1 >= len(lines) or "|" not in lines[index]:
        return False
    separator = parse_table_row(lines[index + 1])
    header = parse_table_row(lines[index])
    return bool(header) and len(header) == len(separator) and all(TABLE_SEPARATOR_RE.fullmatch(cell.strip()) for cell in separator)


def parse_table_row(line: str) -> list[str]:
    value = line.strip().strip("|")
    cells: list[str] = []
    current: list[str] = []
    escaped = False
    math_delimiter = ""
    for character in value:
        if escaped:
            current.append(character)
            escaped = False
        elif character == "\\":
            escaped = True
            current.append(character)
        elif character == "$":
            math_delimiter = "" if math_delimiter else "$"
            current.append(character)
        elif character == "|" and not math_delimiter:
            cells.append("".join(current).strip().replace(r"\|", "|"))
            current = []
        else:
            current.append(character)
    cells.append("".join(current).strip().replace(r"\|", "|"))
    return cells


def starts_block(lines: list[str], index: int) -> bool:
    stripped = lines[index].strip()
    return bool(
        re.match(r"^#{1,9}\s+", stripped)
        or LIST_RE.match(lines[index])
        or IMAGE_RE.match(stripped)
        or stripped.startswith((">", "```", "~~~", "$$", r"\["))
        or stripped in {"---", "***", "___", "<!-- pagebreak -->", "\\pagebreak"}
        or is_table_start(lines, index)
    )


def add_styled_run(paragraph: Any, text: str, style: dict[str, Any], theme: dict[str, Any]) -> Any:
    run = paragraph.add_run(text)
    apply_font(run.font, style, theme)
    return run


def add_hyperlink(
    paragraph: Any,
    label: str,
    url: str,
    style: dict[str, Any],
    theme: dict[str, Any],
) -> None:
    relationship = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run_element = OxmlElement("w:r")
    text_element = OxmlElement("w:t")
    text_element.text = label
    run_element.append(text_element)
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)
    from docx.text.run import Run

    run = Run(run_element, paragraph)
    link_style = dict(style)
    link_style["color"] = "0563C1"
    apply_font(run.font, link_style, theme)
    run.font.underline = True


def add_horizontal_rule(doc: Any) -> None:
    paragraph = doc.add_paragraph()
    properties = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "B8C2CC")
    borders.append(bottom)
    properties.append(borders)


def find_unescaped(text: str, delimiter: str, start: int) -> int:
    position = start
    while True:
        position = text.find(delimiter, position)
        if position < 0:
            return -1
        backslashes = 0
        cursor = position - 1
        while cursor >= 0 and text[cursor] == "\\":
            backslashes += 1
            cursor -= 1
        if backslashes % 2 == 0:
            return position
        position += len(delimiter)


def parse_destination(raw: str) -> str:
    value = raw.strip()
    if value.startswith("<") and ">" in value:
        return value[1 : value.index(">")]
    match = re.match(r"^(.*?)(?:\s+[\"'].*[\"'])?$", value)
    return (match.group(1) if match else value).strip()


def empty_stats() -> dict[str, int]:
    return {
        "headings": 0,
        "paragraphs": 0,
        "list_items": 0,
        "tables": 0,
        "images": 0,
        "equations": 0,
        "equation_fallbacks": 0,
        "quotes": 0,
        "code_blocks": 0,
        "page_breaks": 0,
    }
