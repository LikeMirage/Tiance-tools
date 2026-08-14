from __future__ import annotations

from datetime import datetime
import hashlib
import json
import os
from pathlib import Path
import shutil
import tempfile
from typing import Any

from docx import Document

from tiance_runtime import run_tool
from word_errors import WordOperationError
from word_editing import apply_operations, inspect_document
from word_selection import resolve_selection
from word_elements import (
    add_elements,
    apply_core_properties,
    apply_document_defaults,
    apply_page_settings,
    merged_theme,
    set_header_footer,
)


TOOL_BACKUP_ROOT = Path(".Tiance") / "tool-backups" / "word_document_editor"


class ToolError(Exception):
    def __init__(self, code: str, message: str, details: dict[str, Any] | None = None) -> None:
        super().__init__(message)
        self.code = code
        self.message = message
        self.details = details or {}


def ok(summary: str, data: dict[str, Any], warnings: list[str] | None = None) -> dict[str, Any]:
    return {"ok": True, "summary": summary, "data": data, "warnings": warnings or []}


def fail(
    code: str,
    message: str,
    details: dict[str, Any] | None = None,
    warnings: list[str] | None = None,
) -> dict[str, Any]:
    return {
        "ok": False,
        "error": f"{code}: {message}",
        "error_info": {"code": code, "message": message, "details": details or {}},
        "warnings": warnings or [],
    }


def workspace_root() -> Path:
    raw = os.environ.get("TIANCE_WORKSPACE_ROOT") or os.environ.get("WORKSPACE_ROOT") or os.getcwd()
    return Path(raw).expanduser().resolve(strict=False)


def read_bool(value: Any, default: bool = False) -> bool:
    if isinstance(value, bool):
        return value
    if isinstance(value, str):
        lowered = value.strip().lower()
        if lowered in {"true", "1", "yes", "y", "on"}:
            return True
        if lowered in {"false", "0", "no", "n", "off"}:
            return False
    return default


def read_int(value: Any, default: int, *, minimum: int, maximum: int) -> int:
    if not isinstance(value, int):
        value = default
    return max(minimum, min(value, maximum))


def resolve_input_path(value: Any, root: Path) -> Path:
    if not isinstance(value, str) or not value.strip():
        raise ToolError("INVALID_ARGUMENT", "input_path 必须是非空字符串。")
    path = Path(value.strip()).expanduser()
    if not path.is_absolute():
        path = root / path
    resolved = path.resolve(strict=False)
    ensure_inside_workspace(resolved, root, "input_path")
    if not resolved.is_file():
        raise ToolError("INPUT_NOT_FOUND", "输入 DOCX 文件不存在。", {"input_path": str(resolved)})
    if resolved.suffix.lower() != ".docx":
        raise ToolError("UNSUPPORTED_FORMAT", "当前只支持 .docx 文件。", {"input_path": str(resolved)})
    return resolved


def resolve_output_path(value: Any, root: Path) -> Path:
    if not isinstance(value, str) or not value.strip():
        raise ToolError("INVALID_ARGUMENT", "output_path 必须是非空字符串。")
    path = Path(value.strip()).expanduser()
    if not path.is_absolute():
        path = root / path
    resolved = path.resolve(strict=False)
    ensure_inside_workspace(resolved, root, "output_path")
    if resolved.suffix.lower() != ".docx":
        resolved = resolved.with_suffix(".docx")
    return resolved


def ensure_inside_workspace(path: Path, root: Path, field_name: str) -> None:
    try:
        path.relative_to(root)
    except ValueError as exc:
        raise ToolError(
            "PATH_OUTSIDE_WORKSPACE",
            f"{field_name} 必须位于工作区内。",
            {field_name: str(path), "workspace_root": str(root)},
        ) from exc


def ensure_can_write(output_path: Path, *, overwrite: bool) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    if output_path.exists() and not overwrite:
        raise ToolError("OUTPUT_EXISTS", "输出文件已存在。", {"output_path": str(output_path)})


def create_document(payload: dict[str, Any], root: Path) -> dict[str, Any]:
    output_path = resolve_output_path(payload.get("output_path"), root)
    overwrite = read_bool(payload.get("overwrite"), False)
    ensure_can_write(output_path, overwrite=overwrite)

    elements = payload.get("elements")
    if not isinstance(elements, list) or not elements:
        raise ToolError("INVALID_ARGUMENT", "create 操作必须提供非空 elements。")

    theme = merged_theme(payload.get("theme"))
    doc = Document()
    apply_document_defaults(doc, theme)
    apply_page_settings(doc, payload.get("page"))
    apply_core_properties(doc, payload.get("properties"))
    if isinstance(payload.get("header"), str) or isinstance(payload.get("footer"), str):
        set_header_footer(
            doc,
            header=payload.get("header") if isinstance(payload.get("header"), str) else None,
            footer=payload.get("footer") if isinstance(payload.get("footer"), str) else None,
            theme=theme,
        )
    warnings: list[str] = []
    stats = add_elements(doc, elements, theme, root, warnings=warnings)
    save_document_atomic(doc, output_path)
    return ok(
        f"DOCX 创建完成：{output_path.name}。",
        {
            "action": "create",
            "output_path": str(output_path),
            "stats": stats,
            "overwrite": overwrite,
        },
        warnings=warnings,
    )


def inspect_docx(payload: dict[str, Any], root: Path) -> dict[str, Any]:
    input_path = resolve_input_path(payload.get("input_path"), root)
    inspect_options = payload.get("inspect") if isinstance(payload.get("inspect"), dict) else {}
    doc = Document(str(input_path))
    data = inspect_document(
        doc,
        include_paragraphs=read_bool(inspect_options.get("include_paragraphs"), True),
        include_tables=read_bool(inspect_options.get("include_tables"), True),
        max_paragraphs=read_int(inspect_options.get("max_paragraphs"), 80, minimum=1, maximum=500),
        max_text_chars=read_int(inspect_options.get("max_text_chars"), 20000, minimum=1000, maximum=100000),
        max_formulas=read_int(inspect_options.get("max_formulas"), 200, minimum=1, maximum=1000),
        include_formulas=read_bool(
            inspect_options.get("include_formulas"),
            not isinstance(inspect_options.get("selection"), dict),
        ),
    )
    selection_spec = inspect_options.get("selection")
    if isinstance(selection_spec, dict):
        selection_fingerprint = (
            selection_spec.get("document_fingerprint")
            or selection_spec.get("documentFingerprint")
            or inspect_options.get("document_fingerprint")
            or inspect_options.get("documentFingerprint")
        )
        if selection_fingerprint is not None and selection_fingerprint != document_fingerprint(input_path):
            raise ToolError("STALE_REFERENCE", "inspect.selection 使用了过期文档引用，请重新引用。")
        data["selection"] = resolve_selection(doc, selection_spec).summary()
    data.update({"action": "inspect", "input_path": str(input_path), "document_fingerprint": document_fingerprint(input_path)})
    return ok(f"DOCX 检查完成：{input_path.name}。", data)


def edit_document(payload: dict[str, Any], root: Path) -> dict[str, Any]:
    input_path = resolve_input_path(payload.get("input_path"), root)
    output_path = resolve_output_path(payload.get("output_path"), root)
    overwrite = read_bool(payload.get("overwrite"), False)
    backup = read_bool(payload.get("backup"), True)
    dry_run = read_bool(payload.get("dry_run"), False)

    operations = payload.get("operations")
    if not isinstance(operations, list) or not operations:
        raise ToolError("INVALID_ARGUMENT", "edit 操作必须提供非空 operations。")
    if dry_run:
        if output_path.exists() and not overwrite:
            raise ToolError(
                "OUTPUT_EXISTS",
                "输出文件已存在；如需原位编辑或覆盖，请在 dry_run 和正式提交中都设置 overwrite: true。",
                {"output_path": str(output_path), "required_parameter": "overwrite=true"},
            )
    else:
        validation_token = payload.get("validation_token")
        if not isinstance(validation_token, str) or not validation_token:
            raise ToolError("DRY_RUN_REQUIRED", "edit 写入前必须先用完全相同的参数执行 dry_run，并提交返回的 validation_token。")
        ensure_can_write(output_path, overwrite=overwrite)

    input_fingerprint = document_fingerprint(input_path)
    validate_reference_fingerprints(operations, input_fingerprint)
    expected_token = edit_validation_token(
        input_fingerprint,
        output_path,
        operations,
        overwrite=overwrite,
        backup=backup,
    )
    if not dry_run and payload.get("validation_token") != expected_token:
        raise ToolError("STALE_DRY_RUN", "dry_run 令牌已过期，文档或编辑参数发生变化；请重新 dry_run。")
    doc = Document(str(input_path))
    before_stats = {
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "inline_shape_count": len(doc.inline_shapes),
        "section_count": len(doc.sections),
    }
    theme = merged_theme(payload.get("theme"))
    operation_summaries = apply_operations(doc, operations, theme, root)
    after_stats = {
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "inline_shape_count": len(doc.inline_shapes),
        "section_count": len(doc.sections),
    }

    if dry_run:
        return ok(
            "DOCX 编辑预演完成，未写入文件。",
            {
                "action": "edit",
                "dry_run": True,
                "input_path": str(input_path),
                "output_path": str(output_path),
                "before": before_stats,
                "after": after_stats,
                "operations": operation_summaries,
                "document_fingerprint": input_fingerprint,
                "overwrite": overwrite,
                "backup": backup,
                "validation_token": expected_token,
            },
        )

    backup_path = ""
    if input_path == output_path and output_path.exists() and backup:
        backup_path = create_backup(output_path, root)
    save_document_atomic(doc, output_path)
    return ok(
        f"DOCX 编辑完成：{output_path.name}。",
        {
            "action": "edit",
            "input_path": str(input_path),
            "output_path": str(output_path),
            "backup_path": backup_path,
            "before": before_stats,
            "after": after_stats,
            "operations": operation_summaries,
            "overwrite": overwrite,
            "backup": backup,
            "document_fingerprint": input_fingerprint,
        },
    )


def document_fingerprint(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return f"sha256:{digest.hexdigest()}"


def validate_reference_fingerprints(operations: list[Any], current: str) -> None:
    for index, operation in enumerate(operations):
        if not isinstance(operation, dict) or operation.get("type") != "selection":
            continue
        selection = operation.get("selection")
        if not isinstance(selection, dict):
            continue
        expected = selection.get("document_fingerprint") or selection.get("documentFingerprint")
        if expected is not None and expected != current:
            raise ToolError(
                "STALE_REFERENCE",
                f"operations[{index}] 使用了过期文档引用；请重新引用或 inspect。",
                {"expected": expected, "actual": current},
            )


def edit_validation_token(
    input_fingerprint: str,
    output_path: Path,
    operations: list[Any],
    *,
    overwrite: bool,
    backup: bool,
) -> str:
    payload = json.dumps(
        {
            "input": input_fingerprint,
            "output": str(output_path),
            "operations": operations,
            "overwrite": overwrite,
            "backup": backup,
        },
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
    ).encode("utf-8")
    return "dryrun:" + hashlib.sha256(payload).hexdigest()


def create_backup(path: Path, root: Path) -> str:
    relative_path = path.relative_to(root)
    backup_directory = root / TOOL_BACKUP_ROOT / relative_path.parent
    backup_directory.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d%H%M%S")
    backup_path = backup_directory / f"{path.name}.{stamp}.bak"
    collision_index = 2
    while backup_path.exists():
        backup_path = backup_directory / f"{path.name}.{stamp}.{collision_index}.bak"
        collision_index += 1
    shutil.copy2(path, backup_path)
    return str(backup_path)


def save_document_atomic(doc: Any, output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    handle = tempfile.NamedTemporaryFile(
        prefix=f".{output_path.stem}-",
        suffix=".docx.tmp",
        dir=output_path.parent,
        delete=False,
    )
    temporary_path = Path(handle.name)
    handle.close()
    try:
        doc.save(temporary_path)
        os.replace(temporary_path, output_path)
    finally:
        temporary_path.unlink(missing_ok=True)


def run(payload: dict[str, Any]) -> dict[str, Any]:
    try:
        action = str(payload.get("action") or "").strip().lower()
        root = workspace_root()
        if action == "create":
            return create_document(payload, root)
        if action == "inspect":
            return inspect_docx(payload, root)
        if action == "edit":
            return edit_document(payload, root)
        raise ToolError("INVALID_ARGUMENT", "action 必须是 create、inspect 或 edit。", {"action": action})
    except ToolError as exc:
        return fail(exc.code, exc.message, exc.details)
    except WordOperationError as exc:
        return fail(exc.code, exc.message, exc.details)
    except ValueError as exc:
        return fail("INVALID_ARGUMENT", str(exc) or type(exc).__name__)
    except Exception as exc:
        return fail("WORD_TOOL_FAILED", str(exc) or type(exc).__name__)


if __name__ == "__main__":
    run_tool(run)
