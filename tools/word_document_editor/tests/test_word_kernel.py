from __future__ import annotations

import json
import os
from pathlib import Path
import subprocess
import sys
from zipfile import ZipFile

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from lxml import etree

WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
NS = {"w": WORD_NS, "m": MATH_NS}
TOOL_ROOT = Path(__file__).resolve().parents[1]
REPOSITORY_ROOT = Path(__file__).resolve().parents[4]


def call(root: Path, payload: dict[str, object]) -> dict[str, object]:
    if payload.get("action") == "edit" and not payload.get("dry_run") and not payload.get("validation_token"):
        preview_payload = dict(payload)
        preview_payload["dry_run"] = True
        preview = call_raw(root, preview_payload)
        if preview.get("ok") is not True:
            return preview
        payload = dict(payload)
        payload["validation_token"] = preview["data"]["validation_token"]
    return call_raw(root, payload)


def call_raw(root: Path, payload: dict[str, object]) -> dict[str, object]:
    env = os.environ.copy()
    env["TIANCE_WORKSPACE_ROOT"] = str(root)
    env["PYTHONIOENCODING"] = "utf-8"
    entry = str(TOOL_ROOT / "program" / "main.py")
    paths = [
        str(TOOL_ROOT / "program"),
        str(TOOL_ROOT / "dependencies" / "py313" / "site-packages"),
        str(REPOSITORY_ROOT / "runtime" / "python-packages" / "backend" / "py313" / "site-packages"),
        str(REPOSITORY_ROOT / "1_PythonServer"),
    ]
    launcher = (
        "import runpy,sys\n"
        "entry=sys.argv[1]\n"
        "for path in reversed(sys.argv[2:]):\n"
        "    path and path not in sys.path and sys.path.insert(0, path)\n"
        "sys.argv=[entry]\n"
        "runpy.run_path(entry, run_name='__main__')\n"
    )
    completed = subprocess.run(
        [sys.executable, "-c", launcher, entry, *paths],
        input=json.dumps(payload, ensure_ascii=False),
        capture_output=True,
        text=True,
        encoding="utf-8",
        env=env,
        check=False,
        cwd=str(TOOL_ROOT / "program"),
    )
    if not completed.stdout.strip():
        raise AssertionError(f"empty stdout stderr={completed.stderr}")
    return json.loads(completed.stdout)


def document_xml(path: Path) -> etree._Element:
    with ZipFile(path) as package:
        return etree.fromstring(package.read("word/document.xml"))


def create_markdown_docx(markdown: str, output: Path) -> None:
    server_root = str(REPOSITORY_ROOT / "1_PythonServer")
    if server_root not in sys.path:
        sys.path.insert(0, server_root)
    from app.services.document_conversion.markdown_docx.converter import Md2DocxConverter
    from app.services.document_conversion.markdown_docx.word_formatting import FontSettings

    converter = Md2DocxConverter(base_path=output.parent, fonts=FontSettings())
    converter.convert(markdown).save(output)


def test_create_uses_content_aware_table_widths_and_native_formula(tmp_path: Path) -> None:
    output = tmp_path / "result.docx"
    result = call(
        tmp_path,
        {
            "action": "create",
            "output_path": output.name,
            "elements": [
                {
                    "type": "table",
                    "rows": [
                        ["编号", "非常长的工作内容说明"],
                        ["1", "这一列需要明显更多宽度，以减少无意义换行。"],
                    ],
                },
                {"type": "equation", "latex": r"\sum_{i=1}^{N} x_i"},
            ],
        },
    )
    assert result["ok"] is True, result
    root = document_xml(output)
    assert root.xpath("count(.//w:tblLayout[@w:type='fixed'])", namespaces=NS) == 1
    widths = [int(value) for value in root.xpath(".//w:tblGrid[1]/w:gridCol/@w:w", namespaces=NS)]
    assert len(widths) == 2
    assert widths[1] > widths[0]
    assert root.xpath("count(.//m:oMath)", namespaces=NS) >= 1
    assert not list(tmp_path.glob("*.docx.tmp"))


def test_invalid_formula_is_preserved_as_text_with_warning(tmp_path: Path) -> None:
    output = tmp_path / "invalid.docx"
    result = call(
        tmp_path,
        {
            "action": "create",
            "output_path": output.name,
            "elements": [{"type": "equation", "latex": r"\frac{a{b}"}],
        },
    )
    assert result["ok"] is True
    assert any("花括号未配对" in warning for warning in result["warnings"])
    root = document_xml(output)
    assert "\\frac{a{b}" in "".join(root.itertext())


def test_in_place_edit_stores_backup_in_tool_directory(tmp_path: Path) -> None:
    source = tmp_path / "documents" / "source.docx"
    source.parent.mkdir()
    doc = Document()
    doc.add_paragraph("修改前")
    doc.save(source)

    result = call(
        tmp_path,
        {
            "action": "edit",
            "input_path": "documents/source.docx",
            "output_path": "documents/source.docx",
            "overwrite": True,
            "operations": [
                {
                    "type": "replace_text",
                    "old_text": "修改前",
                    "new_text": "修改后",
                }
            ],
        },
    )

    assert result["ok"] is True, result
    backup_path = Path(result["data"]["backup_path"])
    expected_directory = (
        tmp_path / ".Tiance" / "tool-backups" / "word_document_editor" / "documents"
    )
    assert backup_path.parent == expected_directory
    assert backup_path.name.startswith("source.docx.")
    assert backup_path.suffix == ".bak"
    assert backup_path.is_file()
    assert not list(source.parent.glob("*.bak"))
    assert Document(str(backup_path)).paragraphs[0].text == "修改前"
    assert Document(str(source)).paragraphs[0].text == "修改后"


def test_selection_replaces_text_and_removes_inline_equation_without_touching_anchors(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "edited.docx"
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("左边界")
    equation = OxmlElement("m:oMath")
    math_run = OxmlElement("m:r")
    math_text = OxmlElement("m:t")
    math_text.text = "x=1"
    math_run.append(math_text)
    equation.append(math_run)
    paragraph._p.append(equation)
    paragraph.add_run("待替换内容")
    paragraph.add_run("右边界")
    doc.save(source)

    result = call(
        tmp_path,
        {
            "action": "edit",
            "input_path": source.name,
            "output_path": output.name,
            "operations": [
                {
                    "type": "selection",
                    "selection": {"start_anchor": "左边界", "end_anchor": "右边界"},
                    "action": "replace",
                    "content_mode": "text",
                    "content": "新内容",
                    "style": {"bold": True, "color": "C00000"},
                }
            ],
        },
    )
    assert result["ok"] is True, result
    root = document_xml(output)
    assert "".join(root.xpath(".//w:t/text()", namespaces=NS)) == "左边界新内容右边界"
    assert root.xpath("count(.//m:oMath)", namespaces=NS) == 0
    assert root.xpath("count(.//w:r[w:t='新内容']/w:rPr/w:b)", namespaces=NS) == 1
    assert root.xpath(".//w:r[w:t='新内容']/w:rPr/w:color/@w:val", namespaces=NS) == ["C00000"]


def test_zero_width_selection_only_allows_insert(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    inserted = tmp_path / "inserted.docx"
    doc = Document()
    doc.add_paragraph("锚点后文")
    doc.save(source)

    result = call(
        tmp_path,
        {
            "action": "edit",
            "input_path": source.name,
            "output_path": inserted.name,
            "operations": [
                {
                    "type": "selection",
                    "selection": {"start_anchor": "锚点"},
                    "action": "insert",
                    "content": "插入",
                }
            ],
        },
    )
    assert result["ok"] is True, result
    assert "".join(document_xml(inserted).xpath(".//w:t/text()", namespaces=NS)) == "锚点插入后文"

    rejected = call(
        tmp_path,
        {
            "action": "edit",
            "input_path": source.name,
            "output_path": "rejected.docx",
            "operations": [
                {
                    "type": "selection",
                    "selection": {"start_anchor": "锚点"},
                    "action": "delete",
                }
            ],
        },
    )
    assert rejected["ok"] is False
    assert "零长度选区只能执行 insert" in rejected["error"]

    expanded = call(
        tmp_path,
        {
            "action": "edit",
            "input_path": source.name,
            "output_path": "expanded.docx",
            "operations": [
                {
                    "type": "selection",
                    "selection": {"start_anchor": "锚点", "expand": "paragraph_end"},
                    "action": "replace",
                    "content": "新结尾",
                }
            ],
        },
    )
    assert expanded["ok"] is True, expanded
    assert "".join(document_xml(tmp_path / "expanded.docx").xpath(".//w:t/text()", namespaces=NS)) == "锚点新结尾"


def test_inclusive_boundary_replaces_anchors_themselves(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "inclusive.docx"
    doc = Document()
    doc.add_paragraph("开头中间结尾")
    doc.save(source)
    result = call(
        tmp_path,
        {
            "action": "edit",
            "input_path": source.name,
            "output_path": output.name,
            "operations": [
                {
                    "type": "selection",
                    "selection": {
                        "start_anchor": "开头",
                        "end_anchor": "结尾",
                        "boundary_mode": "inclusive",
                    },
                    "action": "replace",
                    "content": "整段新句",
                }
            ],
        },
    )
    assert result["ok"] is True, result
    assert "".join(document_xml(output).xpath(".//w:t/text()", namespaces=NS)) == "整段新句"


def test_formula_anchor_replaces_matching_word_equation(tmp_path: Path) -> None:
    source = tmp_path / "formula.docx"
    output = tmp_path / "formula_edited.docx"
    created = call(
        tmp_path,
        {
            "action": "create",
            "output_path": source.name,
            "elements": [
                {"type": "paragraph", "text": "公式如下："},
                {"type": "equation", "latex": r"\sum_{i=1}^{N} x_i"},
                {"type": "paragraph", "text": "结束"},
            ],
        },
    )
    assert created["ok"] is True, created
    result = call(
        tmp_path,
        {
            "action": "edit",
            "input_path": source.name,
            "output_path": output.name,
            "operations": [
                {
                    "type": "selection",
                    "selection": {"start_anchor": r"$$\sum_{i=1}^{N} x_i$$"},
                    "action": "replace",
                    "content_mode": "markdown",
                    "content": r"$$\sum_{i=1}^{n} i$$",
                }
            ],
        },
    )
    assert result["ok"] is True, result
    assert result["data"]["operations"][0]["selection"]["matched_formulas"]
    root = document_xml(output)
    assert root.xpath("count(.//m:oMath)", namespaces=NS) >= 1
    text = "".join(root.xpath(".//w:t/text()", namespaces=NS))
    assert "公式如下：" in text and "结束" in text


def test_markdown_table_formula_can_be_inspected_and_replaced(tmp_path: Path) -> None:
    source = tmp_path / "markdown_table.docx"
    output = tmp_path / "formula_ref_edited.docx"
    create_markdown_docx(
        "| 名称 | 公式 |\n"
        "| --- | --- |\n"
        r"| 电阻并联 | $\frac{1}{R} = \frac{1}{R_1} + \frac{1}{R_2}$ |",
        source,
    )

    inspected = call(tmp_path, {"action": "inspect", "input_path": source.name})
    assert inspected["ok"] is True, inspected
    formulas = inspected["data"]["formulas"]
    assert len(formulas) == 1
    formula_ref = formulas[0]["formula_ref"]
    assert formula_ref.startswith("formula:") and len(formula_ref) == 24
    assert formulas[0]["location"]["container"].startswith("table:0:1:1:")

    edited = call(
        tmp_path,
        {
            "action": "edit",
            "input_path": source.name,
            "output_path": output.name,
            "operations": [
                {
                    "type": "selection",
                    "selection": {"formula_ref": formula_ref},
                    "action": "replace",
                    "content_mode": "equation",
                    "content": r"i\hbar\frac{\partial}{\partial t}\Psi(\mathbf{r},t)=\hat{H}\Psi(\mathbf{r},t)",
                    "style": {"color": "FF0000", "bold": True},
                }
            ],
        },
    )
    assert edited["ok"] is True, edited
    root = document_xml(output)
    assert root.xpath("count(.//m:oMath)", namespaces=NS) == 1
    formula_text = "".join(root.xpath(".//m:t/text()", namespaces=NS))
    assert "Ψ" in formula_text and "H" in formula_text
    math_runs = root.xpath(".//m:oMath//m:r", namespaces=NS)
    assert math_runs
    assert all(
        run.xpath("./w:rPr/w:color/@w:val", namespaces=NS) == ["FF0000"]
        for run in math_runs
    )
    assert all(run.xpath("count(./w:rPr/w:b)", namespaces=NS) == 1 for run in math_runs)


def test_formula_ref_format_changes_equation_color_without_replacing_content(tmp_path: Path) -> None:
    source = tmp_path / "formula_format.docx"
    output = tmp_path / "formula_format_edited.docx"
    created = call(
        tmp_path,
        {
            "action": "create",
            "output_path": source.name,
            "elements": [{"type": "equation", "latex": r"E=mc^2"}],
        },
    )
    assert created["ok"] is True, created
    inspected = call(tmp_path, {"action": "inspect", "input_path": source.name})
    formula_ref = inspected["data"]["formulas"][0]["formula_ref"]

    edited = call(
        tmp_path,
        {
            "action": "edit",
            "input_path": source.name,
            "output_path": output.name,
            "operations": [
                {
                    "type": "selection",
                    "selection": {"formula_ref": formula_ref},
                    "action": "format",
                    "style": {"color": "FFFF00"},
                }
            ],
        },
    )
    assert edited["ok"] is True, edited
    assert edited["data"]["operations"][0]["formatted"]["equations_formatted"] == 1
    root = document_xml(output)
    assert "E=mc2" == "".join(root.xpath(".//m:t/text()", namespaces=NS))
    math_runs = root.xpath(".//m:oMath//m:r", namespaces=NS)
    assert math_runs
    assert all(
        run.xpath("./w:rPr/w:color/@w:val", namespaces=NS) == ["FFFF00"]
        for run in math_runs
    )


def test_markdown_table_formula_anchor_accepts_markdown_formula_replacement(tmp_path: Path) -> None:
    source = tmp_path / "markdown_table.docx"
    output = tmp_path / "markdown_formula_edited.docx"
    old_formula = r"\frac{1}{R} = \frac{1}{R_1} + \frac{1}{R_2}"
    create_markdown_docx(
        "| 名称 | 公式 |\n| --- | --- |\n" + f"| 电阻并联 | ${old_formula}$ |",
        source,
    )
    edited = call(
        tmp_path,
        {
            "action": "edit",
            "input_path": source.name,
            "output_path": output.name,
            "operations": [
                {
                    "type": "selection",
                    "selection": {"start_anchor": f"${old_formula}$"},
                    "action": "replace",
                    "content_mode": "markdown",
                    "content": r"$$E=mc^2$$",
                }
            ],
        },
    )
    assert edited["ok"] is True, edited
    operation = edited["data"]["operations"][0]
    assert operation["content_mode"] == "equation"
    assert operation["source_content_mode"] == "markdown"
    root = document_xml(output)
    assert root.xpath("count(.//m:oMath)", namespaces=NS) == 1
    assert "E=mc2" == "".join(root.xpath(".//m:t/text()", namespaces=NS))


def test_formula_end_anchor_uses_document_paragraph_order(tmp_path: Path) -> None:
    source = tmp_path / "formula_end.docx"
    created = call(
        tmp_path,
        {
            "action": "create",
            "output_path": source.name,
            "elements": [
                {"type": "paragraph", "text": "前文一"},
                {"type": "paragraph", "text": "前文二"},
                {"type": "paragraph", "text": "起点"},
                {"type": "equation", "latex": r"x^2"},
            ],
        },
    )
    assert created["ok"] is True, created
    preview = call(
        tmp_path,
        {
            "action": "inspect",
            "input_path": source.name,
            "inspect": {
                "selection": {
                    "start_anchor": "起点",
                    "end_anchor": r"$x^2$",
                    "boundary_mode": "inclusive",
                }
            },
        },
    )
    assert preview["ok"] is True, preview
    assert len(preview["data"]["selection"]["matched_formula_refs"]) == 1
    assert preview["data"]["selection"]["matched_formula_refs"][0].startswith("formula:")


def test_formula_match_failure_returns_inspectable_candidate(tmp_path: Path) -> None:
    source = tmp_path / "formula_candidate.docx"
    created = call(
        tmp_path,
        {
            "action": "create",
            "output_path": source.name,
            "elements": [{"type": "equation", "latex": r"x^2+y^2"}],
        },
    )
    assert created["ok"] is True, created
    failed = call(
        tmp_path,
        {
            "action": "inspect",
            "input_path": source.name,
            "inspect": {"selection": {"start_anchor": r"$x^3+y^3$"}},
        },
    )
    assert failed["ok"] is False
    assert failed["error_info"]["code"] == "SELECTION_NOT_FOUND"
    assert "formula:" in failed["error"]
    assert "inspect" in failed["error"]


def test_formula_ref_replaces_adjacent_formula_in_place(tmp_path: Path) -> None:
    source = tmp_path / "adjacent.docx"
    output = tmp_path / "adjacent_edited.docx"
    doc = Document()
    paragraph = doc.add_paragraph()
    for value in ("x", "y"):
        equation = OxmlElement("m:oMath")
        run = OxmlElement("m:r")
        text = OxmlElement("m:t")
        text.text = value
        run.append(text)
        equation.append(run)
        paragraph._p.append(equation)
    doc.save(source)
    inspected = call(tmp_path, {"action": "inspect", "input_path": source.name})
    formula_ref = inspected["data"]["formulas"][1]["formula_ref"]

    edited = call(
        tmp_path,
        {
            "action": "edit",
            "input_path": source.name,
            "output_path": output.name,
            "operations": [
                {
                    "type": "selection",
                    "selection": {"formula_ref": formula_ref},
                    "action": "replace",
                    "content_mode": "equation",
                    "content": "z",
                }
            ],
        },
    )
    assert edited["ok"] is True, edited
    root = document_xml(output)
    assert root.xpath(".//m:oMath/m:r/m:t/text()", namespaces=NS) == ["x", "z"]


def test_same_cell_cross_paragraph_selection_is_supported(tmp_path: Path) -> None:
    source = tmp_path / "table.docx"
    output = tmp_path / "table_multi.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = "第一段"
    cell.add_paragraph("第二段")
    doc.save(source)
    result = call(
        tmp_path,
        {
            "action": "edit",
            "input_path": source.name,
            "output_path": output.name,
            "operations": [
                {
                    "type": "selection",
                    "selection": {
                        "start_anchor": "第一段",
                        "end_anchor": "第二段",
                        "boundary_mode": "inclusive",
                    },
                    "action": "replace",
                    "content": "合并",
                }
            ],
        },
    )
    assert result["ok"] is True, result
    assert "合并" in "".join(document_xml(output).xpath(".//w:t/text()", namespaces=NS))
    assert "第一段" not in "".join(document_xml(output).xpath(".//w:t/text()", namespaces=NS))
    assert "第二段" not in "".join(document_xml(output).xpath(".//w:t/text()", namespaces=NS))


def test_selection_formats_only_existing_selected_text(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "formatted.docx"
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("开始")
    target_run = paragraph.add_run("需要格式化")
    target_run.font.size = Pt(10)
    target_run.font.color.rgb = RGBColor(0x22, 0x33, 0x44)
    paragraph.add_run("结束")
    doc.save(source)

    result = call(
        tmp_path,
        {
            "action": "edit",
            "input_path": source.name,
            "output_path": output.name,
            "operations": [
                {
                    "type": "selection",
                    "selection": {"start_anchor": "开始", "end_anchor": "结束"},
                    "action": "format",
                    "style": {"bold": True, "italic": True, "font_size": 14, "color": "1F4E79"},
                }
            ],
        },
    )
    assert result["ok"] is True, result
    root = document_xml(output)
    assert root.xpath("count(.//w:r[w:t='需要格式化']/w:rPr/w:b)", namespaces=NS) == 1
    assert root.xpath("count(.//w:r[w:t='需要格式化']/w:rPr/w:i)", namespaces=NS) == 1
    assert root.xpath(".//w:r[w:t='需要格式化']/w:rPr/w:color/@w:val", namespaces=NS) == ["1F4E79"]
    assert root.xpath("count(.//w:r[w:t='开始']/w:rPr/w:b)", namespaces=NS) == 0
    assert root.xpath("count(.//w:r[w:t='结束']/w:rPr/w:b)", namespaces=NS) == 0


def test_format_patch_does_not_reset_unspecified_existing_style(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "formatted.docx"
    doc = Document()
    paragraph = doc.add_paragraph("左")
    run = paragraph.add_run("目标")
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(0x12, 0x34, 0x56)
    paragraph.add_run("右")
    doc.save(source)
    result = call(
        tmp_path,
        {
            "action": "edit",
            "input_path": source.name,
            "output_path": output.name,
            "operations": [
                {
                    "type": "selection",
                    "selection": {"start_anchor": "左", "end_anchor": "右"},
                    "action": "format",
                    "style": {"bold": True},
                }
            ],
        },
    )
    assert result["ok"] is True, result
    root = document_xml(output)
    assert root.xpath(".//w:r[w:t='目标']/w:rPr/w:color/@w:val", namespaces=NS) == ["123456"]
    assert root.xpath(".//w:r[w:t='目标']/w:rPr/w:sz/@w:val", namespaces=NS) == ["30"]


def test_cross_block_markdown_replace_removes_old_table_and_writes_native_content(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "markdown.docx"
    doc = Document()
    doc.add_paragraph("左边界")
    doc.add_paragraph("旧段落")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "旧表格"
    doc.add_paragraph("右边界")
    doc.save(source)

    result = call(
        tmp_path,
        {
            "action": "edit",
            "input_path": source.name,
            "output_path": output.name,
            "operations": [
                {
                    "type": "selection",
                    "selection": {"start_anchor": "左边界", "end_anchor": "右边界"},
                    "action": "replace",
                    "content_mode": "markdown",
                    "content": "## 新章节\n\n带有行内公式 $x^2+y^2=z^2$。\n\n| 项目 | 公式 |\n| --- | --- |\n| 面积 | $S=\\pi r^2$ |",
                }
            ],
        },
    )
    assert result["ok"] is True, result
    root = document_xml(output)
    text = "".join(root.xpath(".//w:t/text()", namespaces=NS))
    assert "左边界" in text and "右边界" in text
    assert "旧段落" not in text and "旧表格" not in text
    assert "新章节" in text and "面积" in text
    assert root.xpath("count(.//w:tbl)", namespaces=NS) == 1
    assert root.xpath("count(.//m:oMath)", namespaces=NS) >= 2
    body_paragraphs = [
        "".join(paragraph.xpath(".//w:t/text()", namespaces=NS))
        for paragraph in root.xpath("/w:document/w:body/w:p", namespaces=NS)
    ]
    assert body_paragraphs == ["左边界", "新章节", "带有行内公式 。", "右边界"]


def test_selection_extract_and_dry_run_report_range_without_writing(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "unused.docx"
    doc = Document()
    doc.add_paragraph("左边界需要提取右边界")
    doc.save(source)
    result = call(
        tmp_path,
        {
            "action": "edit",
            "input_path": source.name,
            "output_path": output.name,
            "dry_run": True,
            "operations": [
                {
                    "type": "selection",
                    "selection": {"start_anchor": "左边界", "end_anchor": "右边界"},
                    "action": "extract",
                }
            ],
        },
    )
    assert result["ok"] is True, result
    assert result["data"]["operations"][0]["selection"]["selected_text"] == "需要提取"
    assert not output.exists()

    preview = call(
        tmp_path,
        {
            "action": "inspect",
            "input_path": source.name,
            "inspect": {"selection": {"start_anchor": "左边界", "end_anchor": "右边界"}},
        },
    )
    assert preview["ok"] is True, preview
    assert preview["data"]["selection"]["selected_text"] == "需要提取"


def test_same_table_cell_selection_is_supported_but_cross_cell_range_is_rejected(tmp_path: Path) -> None:
    source = tmp_path / "table.docx"
    output = tmp_path / "table_edited.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "左选区右"
    table.cell(0, 1).text = "另一单元格"
    doc.save(source)
    result = call(
        tmp_path,
        {
            "action": "edit",
            "input_path": source.name,
            "output_path": output.name,
            "operations": [
                {
                    "type": "selection",
                    "selection": {"start_anchor": "左", "end_anchor": "右"},
                    "action": "replace",
                    "content": "新",
                }
            ],
        },
    )
    assert result["ok"] is True, result
    assert "左新右" in "".join(document_xml(output).xpath(".//w:t/text()", namespaces=NS))

    rejected = call(
        tmp_path,
        {
            "action": "edit",
            "input_path": source.name,
            "output_path": "bad.docx",
            "operations": [
                {
                    "type": "selection",
                    "selection": {"start_anchor": "左", "end_anchor": "另一单元格"},
                    "action": "delete",
                }
            ],
        },
    )
    assert rejected["ok"] is False
    assert "不同单元格" in rejected["error"]


def test_word_range_locates_table_formula_and_local_inspect_is_typed(tmp_path: Path) -> None:
    source = tmp_path / "formula.docx"
    output = tmp_path / "formula_colored.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    paragraph = table.cell(0, 0).paragraphs[0]
    equation = OxmlElement("m:oMath")
    math_run = OxmlElement("m:r")
    math_text = OxmlElement("m:t")
    math_text.text = "x=1"
    math_run.append(math_text)
    equation.append(math_run)
    paragraph._p.append(equation)
    doc.save(source)

    initial = call(tmp_path, {"action": "inspect", "input_path": source.name})
    fingerprint = initial["data"]["document_fingerprint"]
    selection = {
        "word_range": {
            "kind": "word_range",
            "start": {
                "container": "table", "tableIndex": 1, "rowIndex": 1,
                "columnIndex": 1, "cellParagraphIndex": 1, "characterOffset": 0,
            },
            "end": {
                "container": "table", "tableIndex": 1, "rowIndex": 1,
                "columnIndex": 1, "cellParagraphIndex": 1, "characterOffset": 3,
            },
        },
        "document_fingerprint": fingerprint,
        "expected_text": "x=1",
    }
    localized = call(
        tmp_path,
        {
            "action": "inspect",
            "input_path": source.name,
            "inspect": {"document_fingerprint": fingerprint, "selection": selection},
        },
    )

    assert localized["ok"] is True, localized
    assert localized["data"]["formulas"] == []
    assert localized["data"]["selection"]["content_kind"] == "equation"
    assert localized["data"]["selection"]["formulas"][0]["formula_text"] == "x=1"
    assert localized["data"]["selection"]["formulas"][0]["formula_ref"].startswith("formula:")

    edited = call(
        tmp_path,
        {
            "action": "edit",
            "input_path": source.name,
            "output_path": output.name,
            "operations": [{
                "type": "selection",
                "selection": selection,
                "action": "format",
                "style": {"color": "FF0000"},
            }],
        },
    )
    assert edited["ok"] is True, edited
    root = document_xml(output)
    assert root.xpath(".//m:oMath//m:r/w:rPr/w:color/@w:val", namespaces=NS) == ["FF0000"]


def test_edit_requires_matching_dry_run_token_and_reference_fingerprint(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("修改前")
    doc.save(source)
    operation = {"type": "replace_text", "old_text": "修改前", "new_text": "修改后"}

    rejected = call_raw(
        tmp_path,
        {"action": "edit", "input_path": source.name, "output_path": output.name, "operations": [operation]},
    )
    assert rejected["error_info"]["code"] == "DRY_RUN_REQUIRED"

    preview = call_raw(
        tmp_path,
        {"action": "edit", "input_path": source.name, "output_path": output.name, "dry_run": True, "operations": [operation]},
    )
    token = preview["data"]["validation_token"]
    doc = Document(str(source))
    doc.add_paragraph("文档已经变化")
    doc.save(source)
    stale = call_raw(
        tmp_path,
        {
            "action": "edit", "input_path": source.name, "output_path": output.name,
            "operations": [operation], "validation_token": token,
        },
    )
    assert stale["error_info"]["code"] == "STALE_DRY_RUN"


def test_word_range_does_not_include_formula_touching_text_boundary(tmp_path: Path) -> None:
    source = tmp_path / "adjacent_range.docx"
    doc = Document()
    paragraph = doc.add_paragraph()
    equation = OxmlElement("m:oMath")
    run = OxmlElement("m:r")
    math_text = OxmlElement("m:t")
    math_text.text = "x"
    run.append(math_text)
    equation.append(run)
    paragraph._p.append(equation)
    paragraph.add_run("tail")
    doc.save(source)
    result = call(
        tmp_path,
        {
            "action": "inspect",
            "input_path": source.name,
            "inspect": {
                "selection": {
                    "word_range": {
                        "kind": "word_range",
                        "start": {"container": "body", "paragraphIndex": 1, "characterOffset": 1},
                        "end": {"container": "body", "paragraphIndex": 1, "characterOffset": 5},
                    },
                    "expected_text": "tail",
                }
            },
        },
    )
    assert result["ok"] is True, result
    assert result["data"]["selection"]["content_kind"] == "text"
    assert result["data"]["selection"]["equation_count"] == 0


def test_body_word_range_uses_preview_global_paragraph_order(tmp_path: Path) -> None:
    source = tmp_path / "body_after_table.docx"
    doc = Document()
    doc.add_table(rows=1, cols=1).cell(0, 0).text = "表格段落"
    doc.add_paragraph("表格后的正文")
    doc.save(source)
    result = call(
        tmp_path,
        {
            "action": "inspect",
            "input_path": source.name,
            "inspect": {
                "selection": {
                    "word_range": {
                        "kind": "word_range",
                        "start": {"container": "body", "paragraphIndex": 2, "characterOffset": 0},
                        "end": {"container": "body", "paragraphIndex": 2, "characterOffset": 6},
                    },
                    "expected_text": "表格后的正文",
                }
            },
        },
    )
    assert result["ok"] is True, result
    assert result["data"]["selection"]["content"] == "表格后的正文"


def test_formula_reference_is_stable_when_unrelated_formula_is_added(tmp_path: Path) -> None:
    source = tmp_path / "stable.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    paragraph = table.cell(0, 0).paragraphs[0]
    equation = OxmlElement("m:oMath")
    run = OxmlElement("m:r")
    text_node = OxmlElement("m:t")
    text_node.text = "target"
    run.append(text_node)
    equation.append(run)
    paragraph._p.append(equation)
    doc.save(source)
    before = call(tmp_path, {"action": "inspect", "input_path": source.name})
    target_ref = before["data"]["formulas"][0]["formula_ref"]

    doc = Document()
    body_equation = OxmlElement("m:oMath")
    body_run = OxmlElement("m:r")
    body_text = OxmlElement("m:t")
    body_text.text = "other"
    body_run.append(body_text)
    body_equation.append(body_run)
    doc.add_paragraph()._p.append(body_equation)
    table = doc.add_table(rows=1, cols=1)
    target_equation = OxmlElement("m:oMath")
    target_run = OxmlElement("m:r")
    target_text = OxmlElement("m:t")
    target_text.text = "target"
    target_run.append(target_text)
    target_equation.append(target_run)
    table.cell(0, 0).paragraphs[0]._p.append(target_equation)
    doc.save(source)
    after = call(tmp_path, {"action": "inspect", "input_path": source.name})

    assert target_ref in {item["formula_ref"] for item in after["data"]["formulas"]}
