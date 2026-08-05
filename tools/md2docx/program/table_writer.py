from __future__ import annotations

from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

import markdown_tables
import word_formatting as formatting
import word_table_layout as table_formatting
from inline_writer import InlineWriter
from table_layout import calculate_column_widths
from text_measurement import FontTextMeasurer
from word_formatting import FontSettings
from word_template_model import ContentStyleProfile
from table_references import ExternalTable


TABLE_FONT_SIZE_POINTS = 10.5


class TableWriter:
    """Parses and writes one Markdown pipe table."""

    def __init__(
        self,
        document,
        fonts: FontSettings,
        inline: InlineWriter,
        *,
        content_style: ContentStyleProfile | None = None,
        table_sample: dict[str, object] | None = None,
    ) -> None:
        self._document = document
        self._fonts = fonts
        self._inline = inline
        self._content_style = content_style
        self._table_sample = table_sample
        self._font_size_points = (
            content_style.run.size_pt
            if content_style is not None and content_style.run.size_pt is not None
            else TABLE_FONT_SIZE_POINTS
        )
        self._measurer = FontTextMeasurer(
            chinese_font=self._fonts.chinese,
            english_font=self._fonts.english,
            math_font=self._fonts.math,
            size_points=self._font_size_points,
        )

    def close(self) -> None:
        self._measurer.close()

    def add(self, lines: list[str], start: int) -> int:
        table_lines, end_index = _collect_table_lines(lines, start)
        headers = markdown_tables.parse_table_row(table_lines[0])
        alignments = markdown_tables.parse_table_alignments(table_lines[1])
        rows = [
            markdown_tables.normalize_row(markdown_tables.parse_table_row(line), len(headers))
            for line in table_lines[2:]
        ]
        table = self._document.add_table(rows=len(rows) + 1, cols=len(headers))
        self._apply_table_profile(table)
        available_width_points = table_formatting.document_available_width_points(self._document)
        column_widths = calculate_column_widths(
            headers,
            rows,
            available_width_points=available_width_points,
            cell_padding_points=table_formatting.CELL_HORIZONTAL_PADDING_POINTS,
            measurer=self._measurer,
        )
        table_formatting.apply_column_widths(
            table,
            column_widths,
            self._document,
        )
        self._apply_table_alignment(table)
        margins = (
            self._table_sample.get("cell_margins_twips")
            if self._table_sample is not None
            else None
        )
        table_formatting.set_cell_margins(
            table,
            margins if isinstance(margins, dict) else None,
        )
        borders = (
            self._table_sample.get("table_borders")
            if self._table_sample is not None
            else None
        )
        table_formatting.set_table_borders(
            table,
            borders if isinstance(borders, dict) else None,
        )
        table_formatting.set_repeat_header(table.rows[0])
        self._write_header(
            table,
            headers,
            alignments,
            column_widths,
            available_width_points,
        )
        self._write_body(
            table,
            rows,
            alignments,
            column_widths,
            available_width_points,
        )
        return end_index

    def add_external(self, external: ExternalTable) -> None:
        """Insert a table loaded from an external workbook reference."""
        rows = external.rows
        if not rows or not rows[0]:
            raise ValueError("table-ref 指向的 Excel 区域为空。")
        column_count = len(rows[0])
        table = self._document.add_table(rows=len(rows), cols=column_count)
        self._apply_table_profile(table)
        total_width = sum(external.column_widths) or float(column_count)
        percentages = [width * 100.0 / total_width for width in external.column_widths]
        table_formatting.apply_column_widths(table, percentages, self._document)
        table_formatting.set_cell_margins(table)
        table_formatting.set_repeat_header(table.rows[0])

        available_width_points = table_formatting.document_available_width_points(self._document)
        for row_index, row_data in enumerate(rows):
            row = table.rows[row_index]
            if row_index < len(external.row_heights) and external.row_heights[row_index]:
                row.height = Pt(float(external.row_heights[row_index]))
                row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            for column, cell_data in enumerate(row_data):
                cell = row.cells[column]
                cell.vertical_alignment = _vertical_alignment(cell_data.vertical)
                paragraph = cell.paragraphs[0]
                formatting.normalize_table_paragraph(paragraph, style=self._content_style)
                self._inline.write(
                    paragraph,
                    cell_data.value,
                    font_size=Pt(cell_data.font_size or self._font_size_points),
                    max_image_width=Pt(max(18.0, available_width_points * percentages[column] / 100.0 - table_formatting.CELL_HORIZONTAL_PADDING_POINTS)),
                )
                _finish_cell(
                    paragraph,
                    self._fonts,
                    self._content_style,
                    cell_data.font_size or self._font_size_points,
                )
                for run in paragraph.runs:
                    run.bold = cell_data.bold
                    run.italic = cell_data.italic
                    if cell_data.font_color:
                        run.font.color.rgb = RGBColor.from_string(cell_data.font_color)
                formatting.apply_alignment(paragraph, cell_data.horizontal or "left")
                if cell_data.fill:
                    table_formatting.set_cell_shading(cell, cell_data.fill)
                _set_cell_borders(cell, cell_data.borders)

        for top, left, bottom, right in external.merges:
            table.cell(top, left).merge(table.cell(bottom, right))

    def _write_header(
        self,
        table,
        headers: list[str],
        alignments: list[str],
        column_widths: list[float],
        available_width_points: float,
    ) -> None:
        for column, header in enumerate(headers):
            cell = table.rows[0].cells[column]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            table_formatting.set_cell_shading(cell, self._header_shading())
            paragraph = cell.paragraphs[0]
            formatting.normalize_table_paragraph(
                paragraph,
                style=self._content_style,
            )
            self._inline.write(
                paragraph,
                header,
                font_size=Pt(self._font_size_points),
                max_image_width=_cell_image_width(
                    column_widths,
                    column,
                    available_width_points,
                ),
            )
            for run in paragraph.runs:
                run.bold = True
            _finish_cell(
                paragraph,
                self._fonts,
                self._content_style,
                self._font_size_points,
            )
            formatting.apply_alignment(
                paragraph,
                markdown_tables.header_alignment(alignments, column),
            )

    def _write_body(
        self,
        table,
        rows: list[list[str]],
        alignments: list[str],
        column_widths: list[float],
        available_width_points: float,
    ) -> None:
        for row_index, row_data in enumerate(rows):
            for column, cell_text in enumerate(row_data):
                cell = table.rows[row_index + 1].cells[column]
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                paragraph = cell.paragraphs[0]
                formatting.normalize_table_paragraph(
                    paragraph,
                    style=self._content_style,
                )
                self._inline.write(
                    paragraph,
                    cell_text,
                    font_size=Pt(self._font_size_points),
                    max_image_width=_cell_image_width(
                        column_widths,
                        column,
                        available_width_points,
                    ),
                )
                _finish_cell(
                    paragraph,
                    self._fonts,
                    self._content_style,
                    self._font_size_points,
                )
                formatting.apply_alignment(
                    paragraph,
                    markdown_tables.body_alignment(alignments, column),
                )

    def _apply_table_profile(self, table) -> None:
        style_name = (
            self._table_sample.get("style_name")
            if self._table_sample is not None
            else None
        )
        try:
            table.style = (
                style_name
                if isinstance(style_name, str) and style_name
                else "Table Grid"
            )
        except KeyError:
            table.style = "Table Grid"

    def _apply_table_alignment(self, table) -> None:
        if self._table_sample is None:
            return
        alignment = self._table_sample.get("alignment")
        table_alignment = {
            "left": WD_TABLE_ALIGNMENT.LEFT,
            "center": WD_TABLE_ALIGNMENT.CENTER,
            "right": WD_TABLE_ALIGNMENT.RIGHT,
        }.get(alignment)
        if table_alignment is not None:
            table.alignment = table_alignment

    def _header_shading(self) -> str:
        if self._table_sample is None:
            return "F2F2F2"
        value = self._table_sample.get("first_cell_shading")
        return value if isinstance(value, str) and value else "F2F2F2"


def _collect_table_lines(lines: list[str], start: int) -> tuple[list[str], int]:
    table_lines: list[str] = []
    index = start
    while index < len(lines) and "|" in lines[index]:
        table_lines.append(lines[index])
        index += 1
    return table_lines, index - 1


def _finish_cell(
    paragraph,
    fonts: FontSettings,
    style: ContentStyleProfile | None,
    size_points: float,
) -> None:
    formatting.style_runs(
        paragraph.runs,
        fonts,
        style=style,
        default_size=Pt(size_points),
    )
    formatting.set_runs_default_size(paragraph.runs, Pt(size_points))


def _cell_image_width(
    column_widths: list[float],
    column: int,
    available_width_points: float,
):
    percentage = column_widths[column] if column < len(column_widths) else 0.0
    column_width = available_width_points * percentage / 100.0
    content_width = max(18.0, column_width - table_formatting.CELL_HORIZONTAL_PADDING_POINTS)
    return Pt(content_width)


def _vertical_alignment(value: str | None):
    return {
        "top": WD_CELL_VERTICAL_ALIGNMENT.TOP,
        "center": WD_CELL_VERTICAL_ALIGNMENT.CENTER,
        "bottom": WD_CELL_VERTICAL_ALIGNMENT.BOTTOM,
    }.get(value, WD_CELL_VERTICAL_ALIGNMENT.CENTER)


def _set_cell_borders(cell, borders: dict[str, tuple[str, str, int]]) -> None:
    if not borders:
        return
    properties = cell._tc.get_or_add_tcPr()
    container = properties.find(qn("w:tcBorders"))
    if container is None:
        container = OxmlElement("w:tcBorders")
        properties.append(container)
    for side, (style, color, size) in borders.items():
        border = container.find(qn(f"w:{side}"))
        if border is None:
            border = OxmlElement(f"w:{side}")
            container.append(border)
        border.set(qn("w:val"), style)
        border.set(qn("w:sz"), str(size))
        border.set(qn("w:color"), color)
