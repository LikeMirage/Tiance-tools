from __future__ import annotations

import hashlib
import re
from urllib.parse import unquote

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt

import markdown_inline
import word_formatting as formatting
from formula_writer import FormulaWriter
from media_writer import MediaWriter
from note_registry import NoteRegistry
from word_formatting import FontSettings
from word_xml import get_or_add_ordered_child


class InlineWriter:
    """Writes Markdown inline tokens into one existing Word paragraph."""

    def __init__(
        self,
        fonts: FontSettings,
        formulas: FormulaWriter,
        media: MediaWriter,
        notes: NoteRegistry,
    ) -> None:
        self._fonts = fonts
        self._formulas = formulas
        self._media = media
        self._notes = notes

    def write(self, paragraph, text: str, *, font_size=None, max_image_width=None) -> None:
        text = markdown_inline.normalize_typographic_double_quotes(text)
        for inline_token in markdown_inline.tokenize_inline(text):
            if inline_token.kind == "plain":
                add_plain_text(paragraph, inline_token.value)
                continue
            self._write_token(
                paragraph,
                inline_token.kind,
                inline_token.raw,
                inline_token.value,
                font_size=font_size,
                max_image_width=max_image_width,
            )

    def _write_token(
        self,
        paragraph,
        token_type: str,
        token: str,
        value: str,
        *,
        font_size=None,
        max_image_width=None,
    ) -> None:
        if token_type == "code":
            run = paragraph.add_run(value)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            return
        if token_type == "math":
            self._formulas.write(
                paragraph,
                value,
                font_size=font_size,
                max_width=max_image_width,
            )
            return
        if token.startswith("[!["):
            add_plain_text(paragraph, token)
            return
        if token.startswith("!["):
            image = markdown_inline.parse_image_token(token)
            if image is None or not self._media.insert_image(
                paragraph,
                image[1],
                max_width=max_image_width or Inches(1.5),
                alt_text=image[0],
            ):
                add_plain_text(paragraph, token)
            return
        if token.startswith("[^"):
            if not self._notes.add_reference(paragraph, token):
                add_plain_text(paragraph, token)
            return
        if token.startswith("["):
            if not _add_link(paragraph, token, self._fonts, font_size=font_size):
                add_plain_text(paragraph, token)
            return
        if token.startswith(("**", "__")):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            run.font.size = font_size
            return
        if token.startswith("~~"):
            run = paragraph.add_run(token[2:-2])
            run.font.strike = True
            run.font.size = font_size
            return
        if token.startswith(r"\("):
            self._formulas.write(
                paragraph,
                token[2:-2],
                font_size=font_size,
                max_width=max_image_width,
            )
            return
        if token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
            run.font.size = font_size
            return
        add_plain_text(paragraph, token)


def _add_link(paragraph, token: str, fonts: FontSettings, *, font_size=None) -> bool:
    link = markdown_inline.parse_link_token(token)
    if link is None:
        return False
    text, url = link
    hyperlink = OxmlElement("w:hyperlink")
    if url.startswith("#"):
        hyperlink.set(qn("w:anchor"), bookmark_name_for_anchor(url[1:]))
    else:
        relation_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
        hyperlink.set(qn("r:id"), relation_id)
    hyperlink.set(qn("w:history"), "true")
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    r_fonts = get_or_add_ordered_child(properties, "w:rFonts")
    r_fonts.set(qn("w:ascii"), fonts.english)
    r_fonts.set(qn("w:hAnsi"), fonts.english)
    r_fonts.set(qn("w:cs"), fonts.english)
    r_fonts.set(qn("w:eastAsia"), fonts.chinese)
    color = get_or_add_ordered_child(properties, "w:color")
    color.set(qn("w:val"), "0563C1")
    if font_size is not None:
        half_points = str(max(1, round(font_size.pt * 2)))
        get_or_add_ordered_child(properties, "w:sz").set(qn("w:val"), half_points)
        get_or_add_ordered_child(properties, "w:szCs").set(qn("w:val"), half_points)
    underline = get_or_add_ordered_child(properties, "w:u")
    underline.set(qn("w:val"), "single")
    run.append(properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return True


def markdown_heading_anchor(text: str) -> str:
    value = re.sub(r"!\[([^\]]*)\]\([^)]*\)", r"\1", text)
    value = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", value)
    value = re.sub(r"`+([^`]+)`+", r"\1", value)
    value = markdown_inline.strip_html_tags(value)
    value = re.sub(r"[*_~]", "", value).strip().casefold()
    value = re.sub(r"\s+", "-", value)
    value = re.sub(r"[^\w\u3400-\u9fff-]", "", value)
    value = re.sub(r"-+", "-", value).strip("-")
    return value or "section"


def bookmark_name_for_anchor(anchor: str) -> str:
    normalized = unquote(anchor).strip().casefold()
    digest = hashlib.sha1(normalized.encode("utf-8")).hexdigest()[:20]
    return f"md_{digest}"


def add_plain_text(paragraph, text: str) -> None:
    if not text:
        return
    cleaned = re.sub(r"<br\s*/?>", "\n", text, flags=re.IGNORECASE)
    cleaned = markdown_inline.unescape_markdown(cleaned)
    parts = markdown_inline.strip_html_tags(cleaned).split("\n")
    for index, part in enumerate(parts):
        if index > 0:
            paragraph.add_run().add_break()
        if part:
            formatting.add_plain_runs(paragraph, part)
