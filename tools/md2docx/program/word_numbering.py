from __future__ import annotations

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

import word_formatting as formatting
from inline_writer import InlineWriter
from word_formatting import FontSettings
from word_template_model import ContentStyleProfile
from word_xml import get_or_add_ordered_child, insert_ordered_child


MAX_LIST_LEVELS = 9


class NativeListWriter:
    """Writes editable Word numbering while keeping task-list markers explicit."""

    def __init__(
        self,
        document,
        fonts: FontSettings,
        inline: InlineWriter,
        *,
        content_style: ContentStyleProfile | None = None,
    ) -> None:
        self._document = document
        self._fonts = fonts
        self._inline = inline
        self._content_style = content_style
        self._font_size = Pt(
            content_style.run.size_pt
            if content_style is not None and content_style.run.size_pt is not None
            else 12
        )
        self._numbering = document.part.numbering_part.element
        self._abstract_ids: dict[bool, int] = {}
        self._active_levels: dict[int, tuple[bool, int]] = {}

    def add(self, item: dict[str, object]) -> None:
        if item["task"] is not None:
            self.end_sequence()
            self._add_task_item(item)
            return
        ordered = bool(item["ordered"])
        level = max(0, min(MAX_LIST_LEVELS - 1, int(item["level"])))
        start = int(item.get("start", 1))
        for deeper_level in [value for value in self._active_levels if value > level]:
            del self._active_levels[deeper_level]
        active = self._active_levels.get(level)
        if active is None or active[0] != ordered:
            parent = self._closest_parent(level, ordered)
            num_id = (
                parent[1]
                if parent is not None
                else self._append_numbering_instance(
                    ordered=ordered,
                    level=level,
                    start=start,
                )
            )
            if parent is not None and ordered and start != 1:
                self._ensure_start_override(num_id, level, start)
            active = (ordered, num_id)
            self._active_levels[level] = active
        paragraph = self._document.add_paragraph()
        formatting.normalize_paragraph(paragraph, style=self._content_style)
        _apply_numbering(
            paragraph,
            num_id=active[1],
            level=level,
        )
        self._inline.write(paragraph, str(item["text"]), font_size=self._font_size)
        formatting.style_runs(
            paragraph.runs,
            self._fonts,
            style=self._content_style,
            default_size=self._font_size,
        )

    def end_sequence(self) -> None:
        self._active_levels.clear()

    def _closest_parent(self, level: int, ordered: bool) -> tuple[bool, int] | None:
        for parent_level in range(level - 1, -1, -1):
            parent = self._active_levels.get(parent_level)
            if parent is not None and parent[0] == ordered:
                return parent
        return None

    def _add_task_item(self, item: dict[str, object]) -> None:
        paragraph = self._document.add_paragraph()
        formatting.normalize_paragraph(paragraph, style=self._content_style)
        level = int(item["level"])
        text_indent = 0.56 + 0.28 * level
        paragraph.paragraph_format.left_indent = Pt(text_indent * 72)
        paragraph.paragraph_format.first_line_indent = Pt(-0.28 * 72)
        marker = "☑" if item["task"] == "checked" else "☐"
        run = paragraph.add_run(f"{marker}\t")
        run.font.name = "Segoe UI Symbol"
        self._inline.write(paragraph, str(item["text"]), font_size=self._font_size)
        formatting.style_runs(
            paragraph.runs,
            self._fonts,
            style=self._content_style,
            default_size=self._font_size,
        )

    def _append_abstract_numbering(self, *, ordered: bool) -> int:
        abstract_id = _next_numeric_attribute(
            self._numbering,
            "w:abstractNum",
            "w:abstractNumId",
        )
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi_level = OxmlElement("w:multiLevelType")
        multi_level.set(qn("w:val"), "multilevel" if ordered else "hybridMultilevel")
        abstract.append(multi_level)
        for level in range(MAX_LIST_LEVELS):
            abstract.append(_build_level(level, ordered=ordered))
        insert_ordered_child(self._numbering, abstract)
        return abstract_id

    def _append_numbering_instance(self, *, ordered: bool, level: int, start: int) -> int:
        num_id = _next_numeric_attribute(self._numbering, "w:num", "w:numId")
        number = OxmlElement("w:num")
        number.set(qn("w:numId"), str(num_id))
        abstract_id = OxmlElement("w:abstractNumId")
        abstract_id.set(qn("w:val"), str(self._abstract_id(ordered)))
        number.append(abstract_id)
        if ordered and start != 1:
            override = OxmlElement("w:lvlOverride")
            override.set(qn("w:ilvl"), str(max(0, min(MAX_LIST_LEVELS - 1, level))))
            start_override = OxmlElement("w:startOverride")
            start_override.set(qn("w:val"), str(max(1, start)))
            override.append(start_override)
            number.append(override)
        insert_ordered_child(self._numbering, number)
        return num_id

    def _ensure_start_override(self, num_id: int, level: int, start: int) -> None:
        number = next(
            child
            for child in self._numbering.findall(qn("w:num"))
            if child.get(qn("w:numId")) == str(num_id)
        )
        for override in number.findall(qn("w:lvlOverride")):
            if override.get(qn("w:ilvl")) == str(level):
                start_override = override.find(qn("w:startOverride"))
                if start_override is None:
                    start_override = OxmlElement("w:startOverride")
                    override.insert(0, start_override)
                start_override.set(qn("w:val"), str(max(1, start)))
                return
        override = OxmlElement("w:lvlOverride")
        override.set(qn("w:ilvl"), str(level))
        start_override = OxmlElement("w:startOverride")
        start_override.set(qn("w:val"), str(max(1, start)))
        override.append(start_override)
        number.append(override)

    def _abstract_id(self, ordered: bool) -> int:
        abstract_id = self._abstract_ids.get(ordered)
        if abstract_id is None:
            abstract_id = self._append_abstract_numbering(ordered=ordered)
            self._abstract_ids[ordered] = abstract_id
        return abstract_id


def _build_level(level: int, *, ordered: bool):
    level_element = OxmlElement("w:lvl")
    level_element.set(qn("w:ilvl"), str(level))
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level_element.append(start)
    number_format = OxmlElement("w:numFmt")
    number_format.set(qn("w:val"), "decimal" if ordered else "bullet")
    level_element.append(number_format)
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    level_element.append(suffix)
    level_text = OxmlElement("w:lvlText")
    level_text.set(
        qn("w:val"),
        f"%{level + 1}." if ordered else ("•", "◦", "▪")[level % 3],
    )
    level_element.append(level_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level_element.append(justification)
    paragraph_properties = OxmlElement("w:pPr")
    tabs = get_or_add_ordered_child(paragraph_properties, "w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    position = 720 + level * 360
    tab.set(qn("w:pos"), str(position))
    tabs.append(tab)
    indent = get_or_add_ordered_child(paragraph_properties, "w:ind")
    indent.set(qn("w:left"), str(position))
    indent.set(qn("w:hanging"), "360")
    level_element.append(paragraph_properties)
    if not ordered:
        run_properties = OxmlElement("w:rPr")
        fonts = get_or_add_ordered_child(run_properties, "w:rFonts")
        fonts.set(qn("w:ascii"), "Segoe UI Symbol")
        fonts.set(qn("w:hAnsi"), "Segoe UI Symbol")
        level_element.append(run_properties)
    return level_element


def _apply_numbering(paragraph, *, num_id: int, level: int) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    numbering = get_or_add_ordered_child(paragraph_properties, "w:numPr")
    for child in list(numbering):
        numbering.remove(child)
    level_element = OxmlElement("w:ilvl")
    level_element.set(qn("w:val"), str(max(0, min(MAX_LIST_LEVELS - 1, level))))
    numbering.append(level_element)
    number_id = OxmlElement("w:numId")
    number_id.set(qn("w:val"), str(num_id))
    numbering.append(number_id)


def _next_numeric_attribute(parent, child_tag: str, attribute: str) -> int:
    values: list[int] = []
    for child in parent.findall(qn(child_tag)):
        try:
            values.append(int(child.get(qn(attribute), "")))
        except ValueError:
            continue
    return max(values, default=0) + 1
