from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx.enum.text import WD_COLOR_INDEX, WD_PARAGRAPH_ALIGNMENT
from docx.image.image import Image
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from word_table_layout import document_available_width_twips
from word_template_model import ContentStyleProfile, ParagraphStyleProfile, RunStyleProfile
from word_xml import get_or_add_ordered_child

OMML_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
OMML = f"{{{OMML_NS}}}"
DEFAULT_CHINESE_FONT = "微软雅黑"
DEFAULT_ENGLISH_FONT = "Times New Roman"
DEFAULT_MATH_FONT = "Cambria Math"
EMOJI_FONT = "Segoe UI Emoji"
EMOJI_CLUSTER_RE = re.compile(
    r"(?:[\U0001F1E6-\U0001F1FF]{2}|"
    r"[\U0001F300-\U0001FAFF\u2600-\u27BF][\ufe0f\U0001F3FB-\U0001F3FF]?"
    r"(?:\u200d[\U0001F300-\U0001FAFF\u2600-\u27BF][\ufe0f\U0001F3FB-\U0001F3FF]?)*"
    r")"
)


class FontSettings:
    def __init__(
        self,
        *,
        chinese: str = DEFAULT_CHINESE_FONT,
        english: str = DEFAULT_ENGLISH_FONT,
        math: str = DEFAULT_MATH_FONT,
    ) -> None:
        self.chinese = chinese
        self.english = english
        self.math = math


def block_image_max_width(document):
    return Inches(document_available_width_twips(document) / 1440)


def block_image_max_height(document):
    section = document.sections[-1]
    available = int(section.page_height) - int(section.top_margin) - int(section.bottom_margin)
    return max(Inches(1), int(available * 0.9))


def image_render_width(image_path: Path, max_width):
    try:
        image = Image.from_file(str(image_path))
        if image.width and int(image.width) > 0:
            return min(image.width, max_width)
    except Exception:
        pass
    return max_width


def image_render_dimensions(image_path: Path, max_width, max_height=None):
    try:
        image = Image.from_file(str(image_path))
        width = int(image.width)
        height = int(image.height)
        if width > 0 and height > 0:
            scale = min(1.0, int(max_width) / width)
            if max_height is not None:
                scale = min(scale, int(max_height) / height)
            return max(1, round(width * scale)), max(1, round(height * scale))
    except Exception:
        pass
    return max_width, None


def formula_image_render_width(image_path: Path, max_width):
    high_density_width = image_render_width(image_path, max_width * 2)
    return min(max_width, max(Inches(0.25), high_density_width // 2))


def apply_run_fonts(run, fonts: FontSettings) -> None:
    if _is_emoji_run(run):
        return
    run.font.name = fonts.english
    r_fonts = run_r_fonts(run)
    r_fonts.set(qn("w:ascii"), fonts.english)
    r_fonts.set(qn("w:hAnsi"), fonts.english)
    r_fonts.set(qn("w:cs"), fonts.english)
    r_fonts.set(qn("w:eastAsia"), fonts.chinese)


def apply_emoji_run_font(run) -> None:
    run.font.name = EMOJI_FONT
    r_fonts = run_r_fonts(run)
    r_fonts.set(qn("w:ascii"), EMOJI_FONT)
    r_fonts.set(qn("w:hAnsi"), EMOJI_FONT)
    r_fonts.set(qn("w:cs"), EMOJI_FONT)
    r_fonts.set(qn("w:eastAsia"), EMOJI_FONT)


def apply_math_run_font(run, fonts: FontSettings) -> None:
    run.font.name = fonts.math
    r_fonts = run_r_fonts(run)
    r_fonts.set(qn("w:ascii"), fonts.math)
    r_fonts.set(qn("w:hAnsi"), fonts.math)
    r_fonts.set(qn("w:cs"), fonts.math)
    r_fonts.set(qn("w:eastAsia"), fonts.math)


def set_style_fonts(style, fonts: FontSettings) -> None:
    properties = style._element.get_or_add_rPr()
    r_fonts = get_or_add_ordered_child(properties, "w:rFonts")
    r_fonts.set(qn("w:ascii"), fonts.english)
    r_fonts.set(qn("w:hAnsi"), fonts.english)
    r_fonts.set(qn("w:cs"), fonts.english)
    r_fonts.set(qn("w:eastAsia"), fonts.chinese)


def run_r_fonts(run):
    properties = run._element.get_or_add_rPr()
    return get_or_add_ordered_child(properties, "w:rFonts")


def apply_omml_font(root: Any, fonts: FontSettings, size=None) -> None:
    for math_run in root.iter(f"{OMML}r"):
        properties = math_run.find(qn("w:rPr"))
        if properties is None:
            properties = OxmlElement("w:rPr")
            math_properties = math_run.find(f"{OMML}rPr")
            math_run.insert(1 if math_properties is not None else 0, properties)
        r_fonts = get_or_add_ordered_child(properties, "w:rFonts")
        r_fonts.set(qn("w:ascii"), fonts.math)
        r_fonts.set(qn("w:hAnsi"), fonts.math)
        r_fonts.set(qn("w:cs"), fonts.math)
        r_fonts.set(qn("w:eastAsia"), fonts.math)
        if size is not None:
            half_points = str(max(1, round(size.pt * 2)))
            get_or_add_ordered_child(properties, "w:sz").set(qn("w:val"), half_points)
            get_or_add_ordered_child(properties, "w:szCs").set(qn("w:val"), half_points)


def add_plain_runs(paragraph, text: str) -> None:
    position = 0
    for match in EMOJI_CLUSTER_RE.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position:match.start()])
        run = paragraph.add_run(match.group(0))
        apply_emoji_run_font(run)
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])


def normalize_paragraph(
    paragraph,
    *,
    first_line: bool = False,
    style: ContentStyleProfile | None = None,
) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    if first_line:
        paragraph.paragraph_format.first_line_indent = Inches(0.28)
    if style is not None:
        apply_paragraph_profile(paragraph, style.paragraph)


def normalize_heading_paragraph(
    paragraph,
    level: int,
    *,
    style: ContentStyleProfile | None = None,
) -> None:
    paragraph.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
    paragraph.paragraph_format.space_after = Pt(6 if level <= 3 else 4)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.keep_with_next = True
    if style is not None:
        apply_paragraph_profile(paragraph, style.paragraph)


def heading_font_size(level: int, *, style: ContentStyleProfile | None = None):
    if style is not None and style.run.size_pt is not None:
        return Pt(style.run.size_pt)
    if level <= 1:
        return Pt(18)
    if level == 2:
        return Pt(15)
    if level == 3:
        return Pt(13)
    return Pt(12)


def normalize_code_paragraph(paragraph) -> None:
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.right_indent = Inches(0.25)


def normalize_table_paragraph(
    paragraph,
    *,
    style: ContentStyleProfile | None = None,
) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    if style is not None:
        apply_paragraph_profile(paragraph, style.paragraph)


def style_runs(
    runs,
    fonts: FontSettings,
    *,
    style: ContentStyleProfile | None = None,
    default_size=None,
) -> None:
    for run in runs:
        if not run.font.name:
            apply_run_fonts(run, fonts)
        run.font.size = run.font.size or default_size or Pt(12)
        if style is not None:
            apply_run_profile(run, style.run)


def set_runs_default_size(runs, size) -> None:
    for run in runs:
        if run.font.size is None or run.font.size == Pt(12):
            run.font.size = size


def apply_paragraph_profile(paragraph, profile: ParagraphStyleProfile) -> None:
    paragraph_format = paragraph.paragraph_format
    if profile.alignment is not None:
        paragraph_format.alignment = _paragraph_alignment(profile.alignment)
    if profile.space_before_pt is not None:
        paragraph_format.space_before = Pt(profile.space_before_pt)
    if profile.space_after_pt is not None:
        paragraph_format.space_after = Pt(profile.space_after_pt)
    if profile.line_spacing_value is not None:
        paragraph_format.line_spacing = (
            Pt(profile.line_spacing_value)
            if profile.line_spacing_unit == "pt"
            else profile.line_spacing_value
        )
    if profile.left_indent_pt is not None:
        paragraph_format.left_indent = Pt(profile.left_indent_pt)
    if profile.right_indent_pt is not None:
        paragraph_format.right_indent = Pt(profile.right_indent_pt)
    if profile.first_line_indent_pt is not None:
        paragraph_format.first_line_indent = Pt(profile.first_line_indent_pt)
    if profile.keep_together is not None:
        paragraph_format.keep_together = profile.keep_together
    if profile.keep_with_next is not None:
        paragraph_format.keep_with_next = profile.keep_with_next
    if profile.page_break_before is not None:
        paragraph_format.page_break_before = profile.page_break_before
    if profile.widow_control is not None:
        paragraph_format.widow_control = profile.widow_control
    if profile.contextual_spacing is not None:
        properties = paragraph._p.get_or_add_pPr()
        contextual = properties.find(qn("w:contextualSpacing"))
        if contextual is None:
            contextual = OxmlElement("w:contextualSpacing")
            properties.append(contextual)
        contextual.set(qn("w:val"), "1" if profile.contextual_spacing else "0")
    if profile.outline_level is not None:
        properties = paragraph._p.get_or_add_pPr()
        outline = properties.find(qn("w:outlineLvl"))
        if outline is None:
            outline = OxmlElement("w:outlineLvl")
            properties.append(outline)
        outline.set(qn("w:val"), str(profile.outline_level))


def apply_run_profile(run, profile: RunStyleProfile) -> None:
    fonts = run_r_fonts(run)
    if profile.latin_font:
        run.font.name = profile.latin_font
        fonts.set(qn("w:ascii"), profile.latin_font)
        fonts.set(qn("w:hAnsi"), profile.latin_font)
    if profile.complex_script_font:
        fonts.set(qn("w:cs"), profile.complex_script_font)
    if profile.east_asia_font:
        fonts.set(qn("w:eastAsia"), profile.east_asia_font)
    if profile.size_pt is not None:
        run.font.size = Pt(profile.size_pt)
    if profile.bold is not None and run.bold is None:
        run.bold = profile.bold
    if profile.italic is not None and run.italic is None:
        run.italic = profile.italic
    if profile.underline is not None and run.underline is None:
        run.underline = profile.underline
    if profile.strike is not None and run.font.strike is None:
        run.font.strike = profile.strike
    if profile.all_caps is not None and run.font.all_caps is None:
        run.font.all_caps = profile.all_caps
    if profile.small_caps is not None and run.font.small_caps is None:
        run.font.small_caps = profile.small_caps
    if profile.color is not None and run.font.color.rgb is None:
        run.font.color.rgb = RGBColor.from_string(profile.color[-6:])
    if profile.highlight is not None and run.font.highlight_color is None:
        highlight = _highlight_color(profile.highlight)
        if highlight is not None:
            run.font.highlight_color = highlight


def apply_alignment(paragraph, alignment: str) -> None:
    if alignment == "center":
        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif alignment == "right":
        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    else:
        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


def set_shading(paragraph, color: str) -> None:
    shading = get_or_add_ordered_child(paragraph._p.get_or_add_pPr(), "w:shd")
    shading.set(qn("w:fill"), color)


def set_paragraph_border(paragraph, color: str) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    borders = get_or_add_ordered_child(paragraph_properties, "w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        border = borders.find(qn(f"w:{side}"))
        if border is None:
            border = OxmlElement(f"w:{side}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "4")
        border.set(qn("w:color"), color)


def set_left_border(paragraph, color: str) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    borders = get_or_add_ordered_child(paragraph_properties, "w:pBdr")
    border = borders.find(qn("w:left"))
    if border is None:
        border = OxmlElement("w:left")
        borders.append(border)
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), "8")
    border.set(qn("w:space"), "6")
    border.set(qn("w:color"), color)


def set_horizontal_rule(paragraph, color: str = "B7B7B7") -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    borders = get_or_add_ordered_child(paragraph_properties, "w:pBdr")
    border = borders.find(qn("w:bottom"))
    if border is None:
        border = OxmlElement("w:bottom")
        borders.append(border)
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), "6")
    border.set(qn("w:space"), "1")
    border.set(qn("w:color"), color)


def _is_emoji_run(run) -> bool:
    if run.font.name == EMOJI_FONT:
        return True
    text = run.text or ""
    return bool(text and EMOJI_CLUSTER_RE.fullmatch(text))


def _paragraph_alignment(value: str):
    return {
        "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
        "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
        "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
        "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        "distribute": WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE,
        "justify-low": WD_PARAGRAPH_ALIGNMENT.JUSTIFY_LOW,
        "justify-medium": WD_PARAGRAPH_ALIGNMENT.JUSTIFY_MED,
        "justify-high": WD_PARAGRAPH_ALIGNMENT.JUSTIFY_HI,
        "thai-distribute": WD_PARAGRAPH_ALIGNMENT.THAI_JUSTIFY,
    }[value]


def _highlight_color(value: str):
    normalized = value.upper().replace("-", "_")
    return getattr(WD_COLOR_INDEX, normalized, None)
