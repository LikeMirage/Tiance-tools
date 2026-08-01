from __future__ import annotations

from pathlib import Path
from tempfile import TemporaryDirectory

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

import browser_renderer
import word_formatting as formatting
from markdown_inline import parse_image_token
from remote_image import (
    RemoteImageDownloader,
    RemoteImageError,
    is_remote_image_url,
    remote_image_display_url,
)
from warning_collector import WarningCollector
from word_template_model import ContentStyleProfile


class MediaWriter:
    """Owns local image insertion and browser-rendered visual blocks."""

    def __init__(
        self,
        document,
        base_path: Path,
        warnings: WarningCollector,
        render_budget: browser_renderer.BrowserRenderBudget,
        paragraph_style: ContentStyleProfile | None = None,
    ) -> None:
        self._document = document
        self._base_path = base_path
        self._warnings = warnings
        self._render_budget = render_budget
        self._remote_images = RemoteImageDownloader()
        self._paragraph_style = paragraph_style

    def add_mermaid(self, code: str) -> bool:
        try:
            with TemporaryDirectory(prefix="md2docx-mermaid-") as temp_dir:
                image_path = Path(temp_dir) / "mermaid.png"
                browser_renderer.render_mermaid_png(
                    code,
                    image_path,
                    budget=self._render_budget,
                )
                self._add_rendered_image(image_path, alt_text="Mermaid 图")
            return True
        except Exception as exc:
            self._warnings.append(f"Mermaid 渲染失败，已保留源码：{exc}")
            return False

    def add_html(self, html_content: str) -> bool:
        safe_content = browser_renderer.sanitize_html_fragment(html_content)
        try:
            with TemporaryDirectory(prefix="md2docx-html-") as temp_dir:
                image_path = Path(temp_dir) / "html.png"
                browser_renderer.render_html_png(
                    safe_content,
                    image_path,
                    base_path=self._base_path,
                    budget=self._render_budget,
                )
                self._add_rendered_image(image_path, alt_text="HTML 内容截图")
            return True
        except Exception as exc:
            self._warnings.append(f"HTML 块渲染失败，已保留源码：{exc}")
            return False

    def add_block_image(self, token: str) -> None:
        image = parse_image_token(token.strip())
        if image is None:
            return
        paragraph = self._document.add_paragraph()
        formatting.normalize_paragraph(paragraph, style=self._paragraph_style)
        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        if not self.insert_image(
            paragraph,
            image[1],
            max_width=formatting.block_image_max_width(self._document),
            max_height=formatting.block_image_max_height(self._document),
            alt_text=image[0],
        ):
            paragraph.add_run(f"[图片跳过: {image[0] or image[1]}]")

    def insert_image(
        self,
        paragraph,
        image_path: str,
        *,
        max_width,
        max_height=None,
        alt_text: str = "",
    ) -> bool:
        if is_remote_image_url(image_path):
            return self._insert_remote_image(
                paragraph,
                image_path,
                max_width=max_width,
                max_height=max_height,
                alt_text=alt_text,
            )
        resolved = _resolve_local_image_path(image_path, base_path=self._base_path)
        if resolved is None:
            self._warnings.append(f"图片不存在：{image_path}")
            return False
        return self._insert_resolved_image(
            paragraph,
            resolved,
            max_width=max_width,
            max_height=max_height,
            alt_text=alt_text,
            source=image_path,
        )

    def _insert_remote_image(
        self,
        paragraph,
        image_url: str,
        *,
        max_width,
        max_height=None,
        alt_text: str = "",
    ) -> bool:
        display_url = remote_image_display_url(image_url)
        try:
            downloaded = self._remote_images.download(image_url)
            with TemporaryDirectory(prefix="md2docx-remote-image-") as temp_dir:
                image_path = Path(temp_dir) / f"image{downloaded.suffix}"
                image_path.write_bytes(downloaded.content)
                return self._insert_resolved_image(
                    paragraph,
                    image_path,
                    max_width=max_width,
                    max_height=max_height,
                    alt_text=alt_text,
                    source=display_url,
                )
        except RemoteImageError as exc:
            self._warnings.append(f"网络图片下载失败：{display_url}，原因：{exc}")
            return False

    def _insert_resolved_image(
        self,
        paragraph,
        resolved: Path,
        *,
        max_width,
        max_height=None,
        alt_text: str = "",
        source: str,
    ) -> bool:
        try:
            effective_max_height = max_height or formatting.block_image_max_height(self._document)
            width, height = formatting.image_render_dimensions(
                resolved,
                max_width,
                effective_max_height,
            )
            picture_options = {"width": width}
            if height is not None:
                picture_options["height"] = height
            shape = paragraph.add_run().add_picture(str(resolved), **picture_options)
            if alt_text:
                shape._inline.docPr.set("descr", alt_text)
            return True
        except Exception as exc:
            self._warnings.append(f"图片插入失败：{source}，原因：{exc}")
            return False

    def _add_rendered_image(self, image_path: Path, *, alt_text: str) -> None:
        paragraph = self._document.add_paragraph()
        formatting.normalize_paragraph(paragraph, style=self._paragraph_style)
        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        width, height = formatting.image_render_dimensions(
            image_path,
            formatting.block_image_max_width(self._document),
            formatting.block_image_max_height(self._document),
        )
        options = {"width": width}
        if height is not None:
            options["height"] = height
        shape = paragraph.add_run().add_picture(str(image_path), **options)
        shape._inline.docPr.set("descr", alt_text)


def _resolve_local_image_path(image_path: str, *, base_path: Path) -> Path | None:
    raw_path = image_path.strip()
    path = Path(raw_path)
    if not path.is_absolute():
        path = base_path / path
    resolved = path.resolve(strict=False)
    return resolved if resolved.is_file() else None
