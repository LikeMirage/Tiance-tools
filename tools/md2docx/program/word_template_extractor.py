from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable
from uuid import uuid4

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

from word_template_document import (
    extract_document_settings,
    extract_section,
    extract_table_sample,
    extract_theme_fonts,
)
from word_template_model import TEMPLATE_SCHEMA_VERSION, load_template_profile


ROLE_STYLE_IDS = {
    "body": "Normal",
    "title": "Title",
    "subtitle": "Subtitle",
    "quote": "Quote",
    "intense_quote": "IntenseQuote",
    "list": "ListParagraph",
    "caption": "Caption",
    **{f"heading_{level}": f"Heading{level}" for level in range(1, 10)},
}
STYLE_TYPE_NAMES = {
    WD_STYLE_TYPE.PARAGRAPH: "paragraph",
    WD_STYLE_TYPE.CHARACTER: "character",
    WD_STYLE_TYPE.TABLE: "table",
    WD_STYLE_TYPE.LIST: "list",
}


def extract_word_template(
    source_path: Path,
    *,
    template_name: str,
    template_id: str | None = None,
) -> dict[str, Any]:
    document = Document(str(source_path))
    theme_fonts = extract_theme_fonts(document)
    paragraph_styles = {
        style.style_id: style
        for style in document.styles
        if style.type == WD_STYLE_TYPE.PARAGRAPH
    }
    source_styles = [
        _extract_source_style(style, document, theme_fonts)
        for style in document.styles
    ]
    role_styles: dict[str, dict[str, Any]] = {}
    for role, style_id in ROLE_STYLE_IDS.items():
        style = paragraph_styles.get(style_id)
        if style is not None:
            role_styles[role] = _extract_content_style(style, document, theme_fonts)

    if "body" not in role_styles:
        raise ValueError("模板中缺少 Normal 正文样式，无法作为转换模板。")

    table_sample = extract_table_sample(document)
    if table_sample is not None:
        table_style_id = table_sample.get("paragraph_style_id")
        table_style = paragraph_styles.get(str(table_style_id))
        if table_style is not None:
            role_styles["table"] = _extract_content_style(
                table_style,
                document,
                theme_fonts,
            )

    payload: dict[str, Any] = {
        "schema_version": TEMPLATE_SCHEMA_VERSION,
        "template_id": template_id or uuid4().hex,
        "name": template_name,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "source_file_name": source_path.name,
        "source_summary": {
            "sections": len(document.sections),
            "paragraphs": len(document.paragraphs),
            "tables": len(document.tables),
            "styles": len(source_styles),
            "inline_shapes": len(document.inline_shapes),
        },
        "document_settings": extract_document_settings(document),
        "theme_fonts": theme_fonts,
        "sections": [
            extract_section(section, index)
            for index, section in enumerate(document.sections)
        ],
        "role_styles": role_styles,
        "source_styles": source_styles,
        "table_sample": table_sample,
    }
    load_template_profile(payload)
    return payload


def _extract_source_style(
    style,
    document,
    theme_fonts: dict[str, Any],
) -> dict[str, Any]:
    payload: dict[str, Any] = {
        "style_id": style.style_id,
        "name": style.name,
        "type": STYLE_TYPE_NAMES.get(style.type, str(style.type)),
        "builtin": bool(style.builtin),
        "hidden": _optional_bool(getattr(style, "hidden", None)),
        "quick_style": _optional_bool(getattr(style, "quick_style", None)),
        "unhide_when_used": _optional_bool(getattr(style, "unhide_when_used", None)),
        "priority": getattr(style, "priority", None),
        "based_on": _style_id(getattr(style, "base_style", None)),
        "next_style": _style_id(getattr(style, "next_paragraph_style", None)),
        "linked_style": _linked_style_id(style),
        "run": (
            _extract_effective_run_style(style, document, theme_fonts)
            if hasattr(style, "font")
            else None
        ),
    }
    if style.type == WD_STYLE_TYPE.PARAGRAPH:
        payload["paragraph"] = _extract_effective_paragraph_style(style, document)
    return payload


def _extract_content_style(
    style,
    document,
    theme_fonts: dict[str, Any],
) -> dict[str, Any]:
    return {
        "style_id": style.style_id,
        "name": style.name,
        "based_on": _style_id(style.base_style),
        "next_style": _style_id(style.next_paragraph_style),
        "linked_style": _linked_style_id(style),
        "run": _extract_effective_run_style(style, document, theme_fonts),
        "paragraph": _extract_effective_paragraph_style(style, document),
    }


def _extract_effective_run_style(
    style,
    document,
    theme_fonts: dict[str, Any],
) -> dict[str, Any]:
    chain = list(_style_chain(style))
    defaults = _document_default_run_properties(document)
    latin_font = _first_rfont(chain, defaults, ("w:ascii", "w:hAnsi")) or _theme_rfont(
        chain,
        defaults,
        ("w:asciiTheme", "w:hAnsiTheme"),
        theme_fonts,
        script="latin",
    )
    east_asia_font = _first_rfont(chain, defaults, ("w:eastAsia",)) or _theme_rfont(
        chain,
        defaults,
        ("w:eastAsiaTheme",),
        theme_fonts,
        script="east_asia",
    )
    complex_script_font = _first_rfont(chain, defaults, ("w:cs",)) or _theme_rfont(
        chain,
        defaults,
        ("w:cstheme",),
        theme_fonts,
        script="complex_script",
    )
    size_pt = _first_font_value(chain, "size", _length_points)
    if size_pt is None:
        size_half_points = _xml_int(
            defaults.find(qn("w:sz")) if defaults is not None else None,
            "w:val",
        )
        size_pt = size_half_points / 2 if size_half_points is not None else None
    bold = _first_font_value(chain, "bold", _optional_bool)
    italic = _first_font_value(chain, "italic", _optional_bool)
    underline = _first_font_value(chain, "underline", _underline_bool)
    strike = _first_font_value(chain, "strike", _optional_bool)
    all_caps = _first_font_value(chain, "all_caps", _optional_bool)
    small_caps = _first_font_value(chain, "small_caps", _optional_bool)
    return {
        "latin_font": latin_font,
        "east_asia_font": east_asia_font,
        "complex_script_font": complex_script_font,
        "size_pt": size_pt,
        "bold": bold if bold is not None else _xml_on_off(defaults, "w:b"),
        "italic": italic if italic is not None else _xml_on_off(defaults, "w:i"),
        "underline": (
            underline if underline is not None else _xml_on_off(defaults, "w:u")
        ),
        "strike": strike if strike is not None else _xml_on_off(defaults, "w:strike"),
        "all_caps": (
            all_caps if all_caps is not None else _xml_on_off(defaults, "w:caps")
        ),
        "small_caps": (
            small_caps
            if small_caps is not None
            else _xml_on_off(defaults, "w:smallCaps")
        ),
        "color": _first_font_color(chain)
        or _xml_attribute(
            defaults.find(qn("w:color")) if defaults is not None else None,
            "w:val",
        ),
        "highlight": _first_font_value(chain, "highlight_color", _enum_value)
        or _xml_attribute(
            defaults.find(qn("w:highlight")) if defaults is not None else None,
            "w:val",
        ),
    }


def _extract_effective_paragraph_style(style, document) -> dict[str, Any]:
    chain = list(_style_chain(style))
    defaults = _document_default_paragraph_properties(document)
    line_spacing = _first_paragraph_value(chain, "line_spacing")
    line_value, line_unit = _line_spacing_value(line_spacing)
    if line_value is None:
        line_value, line_unit = 1.0, "multiple"
    return {
        "alignment": _enum_value(_first_paragraph_value(chain, "alignment")) or "left",
        "space_before_pt": _length_points(
            _first_paragraph_value(chain, "space_before")
        )
        or 0.0,
        "space_after_pt": _length_points(
            _first_paragraph_value(chain, "space_after")
        )
        or 0.0,
        "line_spacing_value": line_value,
        "line_spacing_unit": line_unit,
        "left_indent_pt": _length_points(
            _first_paragraph_value(chain, "left_indent")
        )
        or 0.0,
        "right_indent_pt": _length_points(
            _first_paragraph_value(chain, "right_indent")
        )
        or 0.0,
        "first_line_indent_pt": _length_points(
            _first_paragraph_value(chain, "first_line_indent")
        )
        or 0.0,
        "keep_together": _optional_bool(
            _first_paragraph_value(chain, "keep_together")
        ),
        "keep_with_next": _optional_bool(
            _first_paragraph_value(chain, "keep_with_next")
        ),
        "page_break_before": _optional_bool(
            _first_paragraph_value(chain, "page_break_before")
        ),
        "widow_control": _optional_bool(
            _first_paragraph_value(chain, "widow_control")
        ),
        "contextual_spacing": _effective_xml_bool(
            chain,
            defaults,
            "w:contextualSpacing",
        ),
        "outline_level": _effective_xml_integer(chain, defaults, "w:outlineLvl"),
    }


def _style_chain(style) -> Iterable[Any]:
    current = style
    visited: set[str] = set()
    while current is not None and current.style_id not in visited:
        visited.add(current.style_id)
        yield current
        current = getattr(current, "base_style", None)


def _document_default_run_properties(document):
    defaults = document.styles.element.find(qn("w:docDefaults"))
    if defaults is None:
        return None
    run_default = defaults.find(qn("w:rPrDefault"))
    return run_default.find(qn("w:rPr")) if run_default is not None else None


def _document_default_paragraph_properties(document):
    defaults = document.styles.element.find(qn("w:docDefaults"))
    if defaults is None:
        return None
    paragraph_default = defaults.find(qn("w:pPrDefault"))
    return (
        paragraph_default.find(qn("w:pPr"))
        if paragraph_default is not None
        else None
    )


def _first_rfont(chain, defaults, attributes: tuple[str, ...]) -> str | None:
    for style in chain:
        properties = style.element.find(qn("w:rPr"))
        value = _rfont_value(properties, attributes)
        if value:
            return value
    return _rfont_value(defaults, attributes)


def _rfont_value(properties, attributes: tuple[str, ...]) -> str | None:
    if properties is None:
        return None
    fonts = properties.find(qn("w:rFonts"))
    if fonts is None:
        return None
    for attribute in attributes:
        value = fonts.get(qn(attribute))
        if value:
            return value
    return None


def _theme_rfont(
    chain,
    defaults,
    attributes: tuple[str, ...],
    theme_fonts: dict[str, Any],
    *,
    script: str,
) -> str | None:
    token = None
    for style in chain:
        properties = style.element.find(qn("w:rPr"))
        token = _rfont_value(properties, attributes)
        if token:
            break
    if not token:
        token = _rfont_value(defaults, attributes)
    if not token:
        return None
    family = "major" if token.lower().startswith("major") else "minor"
    family_fonts = theme_fonts.get(family)
    if not isinstance(family_fonts, dict):
        return None
    value = family_fonts.get(script)
    if isinstance(value, str) and value:
        return value
    if script != "east_asia":
        return None
    supplemental = family_fonts.get("supplemental")
    if not isinstance(supplemental, dict):
        return None
    for language_script in ("Hans", "Hant", "Jpan", "Hang"):
        value = supplemental.get(language_script)
        if isinstance(value, str) and value:
            return value
    return None


def _first_font_value(chain, attribute: str, transform):
    for style in chain:
        value = getattr(style.font, attribute)
        if value is not None:
            return transform(value)
    return None


def _first_font_color(chain) -> str | None:
    for style in chain:
        try:
            value = style.font.color.rgb
        except AttributeError:
            value = None
        if value is not None:
            return str(value).upper()
    return None


def _first_paragraph_value(chain, attribute: str):
    for style in chain:
        value = getattr(style.paragraph_format, attribute)
        if value is not None:
            return value
    return None


def _effective_xml_bool(chain, defaults, tag: str) -> bool | None:
    for style in chain:
        properties = style.element.find(qn("w:pPr"))
        element = properties.find(qn(tag)) if properties is not None else None
        if element is not None:
            return _on_off_value(element)
    element = defaults.find(qn(tag)) if defaults is not None else None
    return _on_off_value(element) if element is not None else None


def _effective_xml_integer(chain, defaults, tag: str) -> int | None:
    for style in chain:
        properties = style.element.find(qn("w:pPr"))
        element = properties.find(qn(tag)) if properties is not None else None
        value = _xml_int(element, "w:val")
        if value is not None:
            return value
    element = defaults.find(qn(tag)) if defaults is not None else None
    return _xml_int(element, "w:val")


def _line_spacing_value(value) -> tuple[float | None, str | None]:
    if value is None:
        return None, None
    if isinstance(value, float):
        return round(value, 4), "multiple"
    points = _length_points(value)
    return points, "pt" if points is not None else None


def _linked_style_id(style) -> str | None:
    link = style.element.find(qn("w:link"))
    return _xml_attribute(link, "w:val")


def _style_id(style) -> str | None:
    return style.style_id if style is not None else None


def _length_points(value) -> float | None:
    return round(float(value.pt), 4) if value is not None else None


def _optional_bool(value) -> bool | None:
    return None if value is None else bool(value)


def _underline_bool(value) -> bool | None:
    if value is None:
        return None
    return bool(value)


def _enum_value(value) -> str | None:
    if value is None:
        return None
    name = getattr(value, "name", None)
    if isinstance(name, str):
        return name.lower().replace("_", "-")
    text = str(value).strip().lower()
    return text.replace("_", "-") or None


def _xml_attribute(element, attribute: str) -> str | None:
    return element.get(qn(attribute)) if element is not None else None


def _xml_int(element, attribute: str, default: int | None = None) -> int | None:
    value = _xml_attribute(element, attribute)
    if value is None:
        return default
    try:
        return int(value)
    except ValueError:
        return default


def _on_off_value(element) -> bool:
    value = element.get(qn("w:val"))
    return value not in {"0", "false", "off", "none"}


def _xml_on_off(properties, tag: str) -> bool | None:
    if properties is None:
        return None
    element = properties.find(qn(tag))
    return _on_off_value(element) if element is not None else None
