from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import browser_renderer
import markdown_blocks
import markdown_preprocessor
import markdown_tables
import word_formatting as formatting
import word_page_layout
from docx_package import save_document_atomically
from formula_writer import FormulaWriter
from inline_writer import (
    InlineWriter,
    bookmark_name_for_anchor,
    markdown_heading_anchor,
)
from markdown_inline import parse_image_token
from media_writer import MediaWriter
from note_registry import NoteRegistry
from table_writer import TableWriter
from word_formatting import FontSettings
from word_numbering import NativeListWriter
from warning_collector import WarningCollector
from word_template_model import ContentStyleProfile, WordTemplateProfile


def convert_markdown_to_docx(
    markdown: str,
    output_path: Path,
    *,
    base_path: Path,
    fonts: FontSettings | None = None,
    page_orientation: str = word_page_layout.DEFAULT_PAGE_ORIENTATION,
    page_size: str = word_page_layout.DEFAULT_PAGE_SIZE,
    template: WordTemplateProfile | None = None,
    overwrite: bool = False,
) -> list[str]:
    converter = Md2DocxConverter(
        base_path=base_path,
        fonts=fonts or FontSettings(),
        page_orientation=page_orientation,
        page_size=page_size,
        template=template,
    )
    document = converter.convert(markdown)
    save_document_atomically(
        document,
        output_path,
        footnotes=converter.footnote_entries,
        endnotes=converter.endnote_entries,
        update_fields=converter.has_toc,
        overwrite=overwrite,
    )
    return converter.warnings


class Md2DocxConverter:
    """Orchestrates Markdown block conversion; specialized writers own inline details."""

    def __init__(
        self,
        *,
        base_path: Path,
        fonts: FontSettings,
        page_orientation: str = word_page_layout.DEFAULT_PAGE_ORIENTATION,
        page_size: str = word_page_layout.DEFAULT_PAGE_SIZE,
        template: WordTemplateProfile | None = None,
    ) -> None:
        self.doc = Document()
        self.fonts = fonts
        self.template = template
        self._warnings = WarningCollector()
        self.has_toc = False
        self._bookmark_id = 0
        self._heading_anchor_counts: dict[str, int] = {}

        self.notes = NoteRegistry(self._warnings)
        render_budget = browser_renderer.BrowserRenderBudget()
        self.formulas = FormulaWriter(self.doc, fonts, self._warnings, render_budget)
        self.media = MediaWriter(
            self.doc,
            base_path,
            self._warnings,
            render_budget,
            paragraph_style=self._style("body"),
        )
        self.inline = InlineWriter(fonts, self.formulas, self.media, self.notes)
        self.tables = TableWriter(
            self.doc,
            fonts,
            self.inline,
            content_style=self._style("table") or self._style("body"),
            table_sample=self.template.table_sample if self.template else None,
        )
        self.lists = NativeListWriter(
            self.doc,
            fonts,
            self.inline,
            content_style=self._style("list") or self._style("body"),
        )
        self._setup_document(page_orientation, page_size)

    @property
    def footnote_entries(self) -> list[tuple[int, str]]:
        return self.notes.footnote_entries

    @property
    def endnote_entries(self) -> list[tuple[int, str]]:
        return self.notes.endnote_entries

    @property
    def warnings(self) -> list[str]:
        return self._warnings.messages()

    def convert(self, content: str) -> Document:
        try:
            return self._convert_content(content)
        finally:
            self.tables.close()

    def _convert_content(self, content: str) -> Document:
        content, invalid_character_count = markdown_preprocessor.sanitize_xml_text(content)
        if invalid_character_count:
            self._warnings.append(
                f"已替换 {invalid_character_count} 个不能写入 Word XML 的控制字符。"
            )
        content, footnotes, endnotes = markdown_preprocessor.prepare_markdown_content(content)
        self.notes.load_definitions(footnotes, endnotes)
        lines = content.split("\n")
        index = 0
        while index < len(lines):
            line = lines[index]
            stripped = line.strip().lstrip("\ufeff")
            if not stripped:
                self.lists.end_sequence()
                index += 1
                continue
            list_item = markdown_blocks.parse_list_item(line)
            if list_item is not None:
                self.lists.add(list_item)
                index += 1
                continue
            self.lists.end_sequence()
            if stripped.upper() in {"[TOC]", "[[TOC]]"}:
                self._add_toc()
                index += 1
                continue
            if stripped.startswith("$$"):
                index = self._add_block_formula(lines, index, opener="$$", closer="$$") + 1
                continue
            if stripped.startswith(r"\["):
                index = self._add_block_formula(lines, index, opener=r"\[", closer=r"\]") + 1
                continue
            if stripped == r"\(":
                index = self._add_block_formula(lines, index, opener=r"\(", closer=r"\)") + 1
                continue
            if re.match(r"^#{1,9}\s+", stripped):
                self._add_heading(stripped)
                index += 1
                continue
            if markdown_tables.is_table_start(lines, index):
                index = self.tables.add(lines, index) + 1
                continue
            if stripped in {"---", "***", "___"}:
                self._add_horizontal_rule()
                index += 1
                continue
            if stripped.startswith(">"):
                index = self._add_blockquote(lines, index) + 1
                continue
            code_fence = markdown_blocks.parse_code_fence(stripped)
            if code_fence is not None:
                index = self._add_code_block(lines, index, code_fence) + 1
                continue
            html_block = markdown_blocks.parse_html_block(lines, index)
            if html_block is not None:
                self._add_html_block(html_block)
                index = html_block["end_index"] + 1
                continue
            if parse_image_token(stripped) is not None:
                self.media.add_block_image(stripped)
                index += 1
                continue
            paragraph_text, index = markdown_blocks.collect_paragraph(lines, index)
            self._add_paragraph(paragraph_text)
        return self.doc

    def _setup_document(self, page_orientation: str, page_size: str) -> None:
        style = self.doc.styles["Normal"]
        style.font.name = self.fonts.english
        formatting.set_style_fonts(style, self.fonts)
        body_style = self._style("body")
        style.font.size = self._font_size(body_style, Pt(12))
        word_page_layout.apply_document_page_layout(
            self.doc,
            page_orientation,
            page_size,
            template_section=self.template.primary_section if self.template else None,
        )

    def _add_block_formula(
        self,
        lines: list[str],
        start: int,
        *,
        opener: str,
        closer: str,
    ) -> int:
        line = lines[start].strip()
        if line.startswith(opener) and line.endswith(closer) and len(line) > len(opener) + len(closer):
            latex = line[len(opener) : -len(closer)].strip()
            end = start
        else:
            latex, end = self._collect_block_formula(lines, start, line, opener, closer)
            if end is None:
                self._warnings.append(f"公式定界符未闭合，已保留原文：{line}")
                paragraph = self.doc.add_paragraph()
                formatting.normalize_paragraph(paragraph, style=self._style("body"))
                formatting.add_plain_runs(paragraph, line)
                return start
        if latex:
            paragraph = self.doc.add_paragraph()
            paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            formatting.normalize_paragraph(paragraph, style=self._style("body"))
            self.formulas.write(
                paragraph,
                latex,
                font_size=self._font_size(self._style("body"), Pt(12)),
                display_mode=True,
            )
        return end

    def _collect_block_formula(
        self,
        lines: list[str],
        start: int,
        line: str,
        opener: str,
        closer: str,
    ) -> tuple[str, int | None]:
        latex_lines: list[str] = []
        index = start + 1 if line == opener else start
        if line != opener:
            latex_lines.append(line[len(opener) :])
            index += 1
        while index < len(lines):
            current = lines[index].strip()
            if current.endswith(closer):
                if current != closer:
                    latex_lines.append(current[: -len(closer)])
                return "\n".join(latex_lines).strip(), index
            latex_lines.append(lines[index])
            index += 1
        return "", None

    def _add_heading(self, line: str) -> None:
        level = min(len(line) - len(line.lstrip("#")), 9)
        heading_text = line[level:].strip()
        paragraph = self.doc.add_heading(level=level)
        heading_style = self._style(f"heading_{level}")
        formatting.normalize_heading_paragraph(paragraph, level, style=heading_style)
        heading_size = formatting.heading_font_size(level, style=heading_style)
        self._add_heading_bookmark(paragraph, heading_text)
        self.inline.write(paragraph, heading_text, font_size=heading_size)
        bookmark_end = OxmlElement("w:bookmarkEnd")
        bookmark_end.set(qn("w:id"), str(self._bookmark_id))
        paragraph._p.append(bookmark_end)
        for run in paragraph.runs:
            formatting.apply_run_fonts(run, self.fonts)
            if heading_style is None:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.size = heading_size
            else:
                formatting.apply_run_profile(run, heading_style.run)

    def _add_paragraph(self, text: str) -> None:
        paragraph = self.doc.add_paragraph()
        body_style = self._style("body")
        body_size = self._font_size(body_style, Pt(12))
        formatting.normalize_paragraph(paragraph, first_line=True, style=body_style)
        self.inline.write(paragraph, text, font_size=body_size)
        formatting.style_runs(
            paragraph.runs,
            self.fonts,
            style=body_style,
            default_size=body_size,
        )

    def _add_heading_bookmark(self, paragraph, heading_text: str) -> None:
        base_anchor = markdown_heading_anchor(heading_text)
        occurrence = self._heading_anchor_counts.get(base_anchor, 0) + 1
        self._heading_anchor_counts[base_anchor] = occurrence
        anchor = base_anchor if occurrence == 1 else f"{base_anchor}-{occurrence}"
        self._bookmark_id += 1
        bookmark = OxmlElement("w:bookmarkStart")
        bookmark.set(qn("w:id"), str(self._bookmark_id))
        bookmark.set(qn("w:name"), bookmark_name_for_anchor(anchor))
        paragraph._p.append(bookmark)

    def _add_horizontal_rule(self) -> None:
        paragraph = self.doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(6)
        formatting.set_horizontal_rule(paragraph)

    def _add_toc(self) -> None:
        self.has_toc = True
        paragraph = self.doc.add_paragraph()
        formatting.normalize_paragraph(paragraph, style=self._style("body"))
        field = OxmlElement("w:fldSimple")
        field.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
        run = OxmlElement("w:r")
        text = OxmlElement("w:t")
        text.text = "目录将在 Word 中自动更新。"
        run.append(text)
        field.append(run)
        paragraph._p.append(field)

    def _add_blockquote(self, lines: list[str], start: int) -> int:
        paragraphs, end_index = markdown_blocks.collect_blockquote(lines, start)
        quote_style = self._style("quote") or self._style("body")
        quote_size = self._font_size(quote_style, Pt(12))
        for text in paragraphs:
            paragraph = self.doc.add_paragraph()
            formatting.normalize_paragraph(paragraph, style=quote_style)
            if self.template is None:
                paragraph.paragraph_format.left_indent = Inches(0.25)
                formatting.set_left_border(paragraph, "A0AEC0")
            self.inline.write(paragraph, text, font_size=quote_size)
            formatting.style_runs(
                paragraph.runs,
                self.fonts,
                style=quote_style,
                default_size=quote_size,
            )
        return end_index

    def _add_code_block(
        self,
        lines: list[str],
        start: int,
        fence: tuple[str, str],
    ) -> int:
        marker, language = fence
        index = start + 1
        code_lines: list[str] = []
        while index < len(lines) and not markdown_blocks.is_code_fence_closer(
            lines[index], marker
        ):
            code_lines.append(lines[index])
            index += 1
        if code_lines:
            code_text = "\n".join(code_lines)
            if language != "mermaid" or not self.media.add_mermaid(code_text):
                self._add_code_source(code_text)
        return index

    def _add_code_source(self, code_text: str) -> None:
        paragraph = self.doc.add_paragraph()
        formatting.normalize_code_paragraph(paragraph)
        formatting.set_shading(paragraph, "F5F5F5")
        formatting.set_paragraph_border(paragraph, "D9D9D9")
        for line_index, code_line in enumerate(code_text.split("\n")):
            if line_index > 0:
                paragraph.add_run().add_break()
            run = paragraph.add_run(code_line or " ")
            run.font.name = "Consolas"
            formatting.run_r_fonts(run).set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(9)

    def _style(self, role: str) -> ContentStyleProfile | None:
        return self.template.style_for(role) if self.template is not None else None

    @staticmethod
    def _font_size(style: ContentStyleProfile | None, default):
        if style is not None and style.run.size_pt is not None:
            return Pt(style.run.size_pt)
        return default

    def _add_html_block(self, html_block: dict[str, Any]) -> None:
        if not html_block["closed"]:
            self._warnings.append(
                f"HTML 块未闭合，已保留源码并继续转换后续内容：<{html_block['root_tag']}>"
            )
            self._add_code_source(html_block["content"])
            return
        if not self.media.add_html(html_block["content"]):
            self._add_code_source(html_block["content"])
