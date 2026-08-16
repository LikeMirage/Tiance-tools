from __future__ import annotations

import hashlib
from pathlib import Path
from tempfile import TemporaryDirectory

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Twips

import browser_renderer
import word_formatting as formatting
from formula_converter import (
    LatexPresentation,
    extract_latex_presentation,
    latex_to_omml,
    preprocess_latex,
    validate_latex,
)
from latex_extensions import needs_image_rendering, normalize_for_image, normalize_for_omml
from word_formatting import FontSettings
from warning_collector import WarningCollector


class FormulaWriter:
    """Converts one LaTeX fragment into editable OMML, an image, or readable source."""

    def __init__(
        self,
        document,
        fonts: FontSettings,
        warnings: WarningCollector,
        render_budget: browser_renderer.BrowserRenderBudget,
    ) -> None:
        self._document = document
        self._fonts = fonts
        self._warnings = warnings
        self._render_budget = render_budget
        self._bookmark_id = 1_000_000
        self._bookmark_names: set[str] = set()

    def write(
        self,
        paragraph,
        latex: str,
        *,
        font_size=None,
        display_mode: bool = False,
        max_width=None,
    ) -> None:
        original = latex.strip()
        if not original:
            return
        presentation, presentation_error = extract_latex_presentation(original)
        if presentation is None:
            self._write_degraded(paragraph, original, presentation_error, font_size=font_size)
            return
        if presentation.tag is not None and not display_mode:
            self._write_degraded(
                paragraph,
                original,
                "\\tag 只能用于独立公式，无法作为 Word 行内公式编号",
                font_size=font_size,
            )
            return
        validation_error = validate_latex(presentation.body)
        if validation_error:
            self._write_degraded(paragraph, original, validation_error, font_size=font_size)
            return
        if needs_image_rendering(presentation.body):
            self._write_image(
                paragraph,
                original,
                presentation,
                display_mode=display_mode,
                max_width=max_width,
                font_size=font_size,
            )
            return
        compatible, notices = normalize_for_omml(presentation.body)
        normalized = preprocess_latex(compatible)
        if not normalized:
            return
        style_notices: list[str] = []
        omml, error = latex_to_omml(normalized, style_notices=style_notices)
        if omml is None:
            self._write_degraded(paragraph, original, error, font_size=font_size)
            return
        bookmark_ids = self._begin_presentation(paragraph, presentation)
        formatting.apply_omml_font(omml, self._fonts, size=font_size)
        paragraph._p.append(omml)
        self._finish_presentation(paragraph, presentation, bookmark_ids, font_size=font_size)
        self._warnings.extend(f"公式兼容转换提示：{notice}" for notice in notices)
        self._warnings.extend(f"公式兼容转换提示：{notice}" for notice in style_notices)

    def _write_image(
        self,
        paragraph,
        original: str,
        presentation: LatexPresentation,
        *,
        display_mode: bool,
        max_width,
        font_size,
    ) -> None:
        try:
            normalized, notices = normalize_for_image(presentation.body)
            with TemporaryDirectory(prefix="md2docx-formula-") as temp_dir:
                image_path = Path(temp_dir) / "formula.png"
                browser_renderer.render_katex_png(
                    normalized,
                    image_path,
                    display_mode=display_mode,
                    budget=self._render_budget,
                )
                picture_run = paragraph.add_run()
                picture_run.add_picture(
                    str(image_path),
                    width=formatting.formula_image_render_width(
                        image_path,
                        max_width or formatting.block_image_max_width(self._document),
                    ),
                )
                bookmark_ids = self._begin_presentation(
                    paragraph,
                    presentation,
                    before=picture_run._r,
                )
                self._finish_presentation(
                    paragraph,
                    presentation,
                    bookmark_ids,
                    font_size=font_size,
                )
            detail = f"；{'；'.join(notices)}" if notices else ""
            self._warnings.append(f"扩展公式已作为图片插入（不可编辑）：{original}{detail}")
        except Exception as exc:
            self._write_degraded(
                paragraph,
                original,
                f"扩展公式图片渲染失败：{exc}",
                font_size=font_size,
            )

    def _begin_presentation(
        self,
        paragraph,
        presentation: LatexPresentation,
        *,
        before=None,
    ) -> list[int]:
        if presentation.tag is not None:
            available = max(1, formatting.document_available_width_twips(self._document))
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            paragraph.paragraph_format.tab_stops.add_tab_stop(
                Twips(available // 2),
                WD_TAB_ALIGNMENT.CENTER,
            )
            paragraph.paragraph_format.tab_stops.add_tab_stop(
                Twips(available),
                WD_TAB_ALIGNMENT.RIGHT,
            )
            tab_run = paragraph.add_run("\t")
            if before is not None:
                before.addprevious(tab_run._r)

        bookmark_ids: list[int] = []
        for label in presentation.labels:
            bookmark_name = self._bookmark_name(label)
            if bookmark_name in self._bookmark_names:
                self._warnings.append(f"公式标签重复，继续引用首次出现的位置：\\label{{{label}}}")
                continue
            self._bookmark_names.add(bookmark_name)
            bookmark_id = self._bookmark_id
            self._bookmark_id += 1
            start = OxmlElement("w:bookmarkStart")
            start.set(qn("w:id"), str(bookmark_id))
            start.set(qn("w:name"), bookmark_name)
            if before is None:
                paragraph._p.append(start)
            else:
                before.addprevious(start)
            bookmark_ids.append(bookmark_id)
        return bookmark_ids

    def _finish_presentation(
        self,
        paragraph,
        presentation: LatexPresentation,
        bookmark_ids: list[int],
        *,
        font_size,
    ) -> None:
        for bookmark_id in reversed(bookmark_ids):
            end = OxmlElement("w:bookmarkEnd")
            end.set(qn("w:id"), str(bookmark_id))
            paragraph._p.append(end)
        if presentation.tag is not None and not presentation.suppress_number:
            run = paragraph.add_run(f"\t({presentation.tag})")
            formatting.apply_math_run_font(run, self._fonts)
            run.font.size = font_size or Pt(10)

    @staticmethod
    def _bookmark_name(label: str) -> str:
        digest = hashlib.sha1(label.encode("utf-8")).hexdigest()[:24]
        return f"tiance_formula_{digest}"

    def _write_degraded(self, paragraph, latex: str, error: str, *, font_size=None) -> None:
        self._warnings.append(f"公式降级为文本：{latex}，原因：{error}")
        run = paragraph.add_run(latex)
        formatting.apply_math_run_font(run, self._fonts)
        run.font.size = font_size or Pt(10)
