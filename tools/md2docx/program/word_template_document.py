from __future__ import annotations

from typing import Any

from docx.oxml import parse_xml
from docx.oxml.ns import qn


def extract_section(section, index: int) -> dict[str, Any]:
    section_properties = section._sectPr
    columns = section_properties.find(qn("w:cols"))
    page_numbering = section_properties.find(qn("w:pgNumType"))
    return {
        "index": index,
        "page_width_mm": _length_mm(section.page_width),
        "page_height_mm": _length_mm(section.page_height),
        "orientation": _orientation(section),
        "top_margin_mm": _length_mm(section.top_margin),
        "bottom_margin_mm": _length_mm(section.bottom_margin),
        "left_margin_mm": _length_mm(section.left_margin),
        "right_margin_mm": _length_mm(section.right_margin),
        "gutter_mm": _length_mm(section.gutter),
        "header_distance_mm": _length_mm(section.header_distance),
        "footer_distance_mm": _length_mm(section.footer_distance),
        "start_type": _enum_value(section.start_type),
        "different_first_page": bool(section.different_first_page_header_footer),
        "vertical_alignment": _xml_attribute(
            section_properties.find(qn("w:vAlign")),
            "w:val",
        ),
        "columns": {
            "count": _xml_int(columns, "w:num", default=1),
            "space_twips": _xml_int(columns, "w:space"),
            "equal_width": _xml_bool(columns, "w:equalWidth", default=True),
            "separator": _xml_bool(columns, "w:sep", default=False),
        },
        "page_numbering": {
            "start": _xml_int(page_numbering, "w:start"),
            "format": _xml_attribute(page_numbering, "w:fmt"),
            "chapter_style": _xml_int(page_numbering, "w:chapStyle"),
            "chapter_separator": _xml_attribute(page_numbering, "w:chapSep"),
        },
        "header_text": _part_paragraph_text(section.header.paragraphs),
        "footer_text": _part_paragraph_text(section.footer.paragraphs),
        "first_page_header_text": _part_paragraph_text(
            section.first_page_header.paragraphs
        ),
        "first_page_footer_text": _part_paragraph_text(
            section.first_page_footer.paragraphs
        ),
        "even_page_header_text": _part_paragraph_text(
            section.even_page_header.paragraphs
        ),
        "even_page_footer_text": _part_paragraph_text(
            section.even_page_footer.paragraphs
        ),
        "header_fields": _part_field_instructions(section.header.paragraphs),
        "footer_fields": _part_field_instructions(section.footer.paragraphs),
    }


def extract_document_settings(document) -> dict[str, Any]:
    settings = document.settings.element
    return {
        "even_and_odd_headers": _child_bool(settings, "w:evenAndOddHeaders"),
        "mirror_margins": _child_bool(settings, "w:mirrorMargins"),
        "book_fold_printing": _child_bool(settings, "w:bookFoldPrinting"),
        "track_revisions": _child_bool(settings, "w:trackRevisions"),
        "auto_hyphenation": _child_bool(settings, "w:autoHyphenation"),
        "do_not_hyphenate_caps": _child_bool(settings, "w:doNotHyphenateCaps"),
        "default_tab_stop_twips": _child_int(settings, "w:defaultTabStop"),
        "consecutive_hyphen_limit": _child_int(settings, "w:consecutiveHyphenLimit"),
        "decimal_symbol": _child_value(settings, "w:decimalSymbol"),
        "list_separator": _child_value(settings, "w:listSeparator"),
        "compatibility_mode": _compatibility_mode(settings),
    }


def extract_theme_fonts(document) -> dict[str, Any]:
    try:
        theme_part = document.part.part_related_by(
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme"
        )
        root = parse_xml(theme_part.blob)
    except (KeyError, AttributeError):
        return {}
    namespace = "http://schemas.openxmlformats.org/drawingml/2006/main"
    result: dict[str, Any] = {}
    for key, xpath in (
        ("major", ".//a:themeElements/a:fontScheme/a:majorFont"),
        ("minor", ".//a:themeElements/a:fontScheme/a:minorFont"),
    ):
        node = root.find(xpath.replace("a:", f"{{{namespace}}}"))
        if node is None:
            continue
        latin = node.find(f"{{{namespace}}}latin")
        east_asia = node.find(f"{{{namespace}}}ea")
        complex_script = node.find(f"{{{namespace}}}cs")
        supplemental = {
            child.get("script"): child.get("typeface")
            for child in node.findall(f"{{{namespace}}}font")
            if child.get("script") and child.get("typeface")
        }
        result[key] = {
            "latin": latin.get("typeface") if latin is not None else None,
            "east_asia": east_asia.get("typeface") if east_asia is not None else None,
            "complex_script": (
                complex_script.get("typeface") if complex_script is not None else None
            ),
            "supplemental": supplemental,
        }
    return result


def extract_table_sample(document) -> dict[str, Any] | None:
    if not document.tables:
        return None
    table = document.tables[0]
    properties = table._tbl.tblPr
    style_element = properties.find(qn("w:tblStyle"))
    margins = properties.find(qn("w:tblCellMar"))
    first_paragraph = table.cell(0, 0).paragraphs[0]
    return {
        "style_id": _xml_attribute(style_element, "w:val"),
        "style_name": table.style.name if table.style is not None else None,
        "alignment": _enum_value(table.alignment),
        "autofit": bool(table.autofit),
        "rows": len(table.rows),
        "columns": len(table.columns),
        "paragraph_style_id": first_paragraph.style.style_id,
        "first_cell_vertical_alignment": _enum_value(table.cell(0, 0).vertical_alignment),
        "cell_margins_twips": {
            side: _table_margin(margins, side)
            for side in ("top", "start", "left", "bottom", "end", "right")
        },
        "first_row_repeats": _row_repeats(table.rows[0]),
        "first_cell_shading": _cell_shading(table.cell(0, 0)),
        "table_borders": _table_borders(properties),
    }


def _orientation(section) -> str:
    value = _enum_value(section.orientation)
    if value in {"portrait", "landscape"}:
        return value
    return "landscape" if section.page_width > section.page_height else "portrait"


def _part_paragraph_text(paragraphs) -> list[str]:
    return [paragraph.text for paragraph in paragraphs if paragraph.text]


def _part_field_instructions(paragraphs) -> list[str]:
    instructions: list[str] = []
    for paragraph in paragraphs:
        for element in paragraph._p.iter(qn("w:instrText")):
            if element.text and element.text.strip():
                instructions.append(element.text.strip())
    return instructions


def _length_mm(value) -> float:
    return round(float(value.mm), 4)


def _enum_value(value) -> str | None:
    if value is None:
        return None
    name = getattr(value, "name", None)
    if isinstance(name, str):
        return name.lower().replace("_", "-")
    text = str(value).strip().lower()
    return text.replace("_", "-") or None


def _xml_attribute(element, attribute: str) -> str | None:
    return element.get(qn(attribute)) if element is not None else None


def _xml_int(element, attribute: str, default: int | None = None) -> int | None:
    value = _xml_attribute(element, attribute)
    if value is None:
        return default
    try:
        return int(value)
    except ValueError:
        return default


def _xml_bool(element, attribute: str, default: bool | None = None) -> bool | None:
    value = _xml_attribute(element, attribute)
    if value is None:
        return default
    return value not in {"0", "false", "off"}


def _on_off_value(element) -> bool:
    value = element.get(qn("w:val"))
    return value not in {"0", "false", "off", "none"}


def _child_bool(parent, tag: str) -> bool:
    element = parent.find(qn(tag))
    return _on_off_value(element) if element is not None else False


def _child_int(parent, tag: str) -> int | None:
    return _xml_int(parent.find(qn(tag)), "w:val")


def _child_value(parent, tag: str) -> str | None:
    return _xml_attribute(parent.find(qn(tag)), "w:val")


def _compatibility_mode(settings) -> int | None:
    compatibility = settings.find(qn("w:compat"))
    if compatibility is None:
        return None
    for setting in compatibility.findall(qn("w:compatSetting")):
        if setting.get(qn("w:name")) == "compatibilityMode":
            try:
                return int(setting.get(qn("w:val"), ""))
            except ValueError:
                return None
    return None


def _table_margin(margins, side: str) -> dict[str, Any] | None:
    if margins is None:
        return None
    element = margins.find(qn(f"w:{side}"))
    if element is None:
        return None
    return {
        "value": _xml_int(element, "w:w"),
        "type": _xml_attribute(element, "w:type"),
    }


def _row_repeats(row) -> bool:
    properties = row._tr.get_or_add_trPr()
    return properties.find(qn("w:tblHeader")) is not None


def _cell_shading(cell) -> str | None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    return _xml_attribute(shading, "w:fill")


def _table_borders(properties) -> dict[str, Any]:
    borders = properties.find(qn("w:tblBorders"))
    if borders is None:
        return {}
    result: dict[str, Any] = {}
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{side}"))
        if element is None:
            continue
        result[side] = {
            "style": _xml_attribute(element, "w:val"),
            "size": _xml_int(element, "w:sz"),
            "space": _xml_int(element, "w:space"),
            "color": _xml_attribute(element, "w:color"),
        }
    return result
