from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

from converter import convert_markdown_to_docx
from word_formatting import FontSettings
from word_template_extractor import extract_word_template
from word_template_store import WordTemplateStore


def test_extract_store_and_apply_word_template(tmp_path: Path) -> None:
    source = tmp_path / "客户样本.docx"
    _build_template_sample(source)

    payload = extract_word_template(source, template_name="客户报告格式")
    assert payload["name"] == "客户报告格式"
    assert payload["source_summary"]["tables"] == 1
    assert payload["sections"][0]["top_margin_mm"] == 25.4
    assert round(payload["sections"][0]["left_margin_mm"], 1) == 31.8
    assert payload["role_styles"]["body"]["run"]["east_asia_font"] == "仿宋"
    assert payload["role_styles"]["body"]["run"]["latin_font"] == "Times New Roman"
    assert payload["role_styles"]["body"]["run"]["size_pt"] == 10.5
    assert payload["role_styles"]["heading_1"]["run"]["size_pt"] == 14.0

    store = WordTemplateStore(tmp_path / "templates")
    profile = store.save(payload)
    listed = store.list_templates()
    assert [item["name"] for item in listed] == ["内置默认样式", "客户报告格式"]
    assert store.load(profile.template_id).name == "客户报告格式"

    output = tmp_path / "result.docx"
    warnings = convert_markdown_to_docx(
        "# 一级标题\n\n这是正文。",
        output,
        base_path=tmp_path,
        fonts=FontSettings(chinese="仿宋", english="Times New Roman"),
        template=profile,
    )
    assert warnings == []

    result = Document(output)
    section = result.sections[0]
    assert round(section.top_margin.mm, 1) == 25.4
    assert round(section.left_margin.mm, 1) == 31.8
    heading = result.paragraphs[0]
    body = result.paragraphs[1]
    assert heading.runs[0].font.size.pt == 14.0
    assert body.runs[0].font.size.pt == 10.5
    assert body.runs[0]._element.rPr.rFonts.get(qn("w:eastAsia")) == "仿宋"
    assert body.paragraph_format.line_spacing == 1.0
    assert body.paragraph_format.space_before.pt == 0.0
    assert body.paragraph_format.space_after.pt == 0.0
    assert section.header.paragraphs[0].text == "客户页眉"
    assert section.footer.paragraphs[0].text == "客户页脚"


def test_template_store_rejects_duplicate_names(tmp_path: Path) -> None:
    source = tmp_path / "样本.docx"
    _build_template_sample(source)
    store = WordTemplateStore(tmp_path / "templates")
    store.save(extract_word_template(source, template_name="同名模板"))

    second = extract_word_template(source, template_name="同名模板")
    try:
        store.save(second)
    except ValueError as exc:
        assert "模板名称已存在" in str(exc)
    else:
        raise AssertionError("重复模板名称应被拒绝")


def _build_template_sample(path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)
    section.left_margin = Mm(31.8)
    section.right_margin = Mm(25.4)
    section.header.paragraphs[0].text = "客户页眉"
    section.footer.paragraphs[0].text = "客户页脚"

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "仿宋")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.first_line_indent = Pt(21)

    for level in (1, 2):
        heading = document.styles[f"Heading {level}"]
        heading.font.name = "Times New Roman"
        heading._element.get_or_add_rPr().get_or_add_rFonts().set(
            qn("w:eastAsia"),
            "仿宋",
        )
        heading.font.size = Pt(14)
        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.space_after = Pt(0)
        heading.paragraph_format.line_spacing = 1.0

    document.add_heading("示例标题", level=1)
    document.add_paragraph("示例正文。")
    table = document.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "表头"
    table.cell(1, 0).text = "内容"
    document.save(path)
