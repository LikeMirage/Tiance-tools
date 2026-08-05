from __future__ import annotations

from pathlib import Path
import sys
from types import SimpleNamespace

from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill

sys.modules.setdefault("tiance_runtime", SimpleNamespace(run_tool=lambda function: None))

from converter import convert_markdown_to_docx
from main import run


def _make_workbook(path: Path) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "汇总"
    sheet.append(["实验结果", None, None])
    sheet.append(["项目", "数值", "备注"])
    sheet.append(["甲", 12, "通过"])
    sheet.merge_cells("A1:C1")
    sheet["A1"].fill = PatternFill("solid", fgColor="D9EAF7")
    sheet["A1"].font = Font(bold=True, color="1F2937")
    sheet["A1"].alignment = Alignment(horizontal="center")
    sheet.column_dimensions["A"].width = 18
    sheet.column_dimensions["B"].width = 12
    sheet.column_dimensions["C"].width = 18
    workbook.save(path)


def test_excel_table_reference_inserts_values_merges_and_styles(tmp_path: Path) -> None:
    workbook_path = tmp_path / "result.xlsx"
    markdown_path = tmp_path / "report.md"
    output_path = tmp_path / "report.docx"
    _make_workbook(workbook_path)
    markdown_path.write_text(
        "# 报告\n\n{{table-ref\nsource=\"result.xlsx\"\nsheet=\"汇总\"\nrange=\"A1:C3\"\n}}\n",
        encoding="utf-8",
    )

    report: list[dict[str, str]] = []
    warnings = convert_markdown_to_docx(
        markdown_path.read_text(encoding="utf-8"),
        output_path,
        base_path=tmp_path,
        table_reference_report=report,
    )

    assert warnings == []
    assert report[0]["sheet"] == "汇总"
    document = Document(output_path)
    assert len(document.tables) == 1
    table = document.tables[0]
    assert table.cell(0, 0).text == "实验结果"
    assert table.cell(1, 1).text == "数值"
    assert table.cell(2, 2).text == "通过"
    assert len(table.cell(0, 0)._tc.xpath("./w:tcPr/w:gridSpan")) == 1


def test_marked_latex_cell_becomes_native_word_equation(tmp_path: Path) -> None:
    workbook_path = tmp_path / "formula.xlsx"
    markdown_path = tmp_path / "formula.md"
    output_path = tmp_path / "formula.docx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["公式", r"$\frac{1}{2}$", "Excel 计算公式"])
    sheet.append(["结果", "1/2", "=SUM(1,2)"])
    workbook.save(workbook_path)
    markdown_path.write_text(
        '{{table-ref source="formula.xlsx" range="A1:C2"}}',
        encoding="utf-8",
    )

    convert_markdown_to_docx(
        markdown_path.read_text(encoding="utf-8"),
        output_path,
        base_path=tmp_path,
    )

    table = Document(output_path).tables[0]
    document_xml = table.cell(0, 1)._tc.xml
    assert "m:oMath" in document_xml
    assert table.cell(1, 2).text == "=SUM(1,2)"


def test_missing_excel_table_reference_fails_without_output(tmp_path: Path) -> None:
    markdown_path = tmp_path / "report.md"
    output_path = tmp_path / "report.docx"
    markdown_path.write_text(
        '{{table-ref source="missing.xlsx" sheet="汇总" range="A1:B2"}}',
        encoding="utf-8",
    )

    result = run({"input_path": str(markdown_path), "output_path": str(output_path)})

    assert result["ok"] is False
    assert "来源文件不存在" in result["error"]
    assert not output_path.exists()
