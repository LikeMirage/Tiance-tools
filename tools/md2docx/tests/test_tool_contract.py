from __future__ import annotations

import json
import sys
from pathlib import Path
from types import SimpleNamespace

import pytest
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

sys.modules.setdefault("tiance_runtime", SimpleNamespace(run_tool=lambda function: None))

import main


TOOL_ROOT = Path(__file__).resolve().parents[1]


def test_tool_success_and_existing_file_contract(tmp_path: Path) -> None:
    input_path = tmp_path / "input.md"
    output_path = tmp_path / "output.docx"
    input_path.write_text("# 契约测试", encoding="utf-8")
    payload = {
        "input_path": str(input_path),
        "output_path": str(output_path),
        "overwrite": False,
    }

    result = main.run(payload)

    assert result == {
        "ok": True,
        "action": "convert",
        "input_path": str(input_path),
        "output_path": str(output_path),
        "overwritten": False,
        "message": "转换完成。",
        "warnings": [],
        "table_references": [],
        "template_id": "builtin-default",
        "template_name": "内置默认样式",
    }
    assert output_path.is_file()

    second = main.run(payload)
    assert second == {
        "ok": False,
        "action": "convert",
        "input_path": str(input_path),
        "output_path": str(output_path),
        "overwritten": False,
        "message": "",
        "warnings": [],
        "table_references": [],
        "error": "目标文件已存在；如需覆盖请设置 overwrite=true。",
    }

    input_path.write_text("# 已更新", encoding="utf-8")
    overwritten = main.run({**payload, "overwrite": True})
    assert overwritten["ok"] is True
    assert overwritten["overwritten"] is True
    assert Document(output_path).paragraphs[0].text == "已更新"


def test_tool_rejects_non_markdown_input_without_creating_output(tmp_path: Path) -> None:
    input_path = tmp_path / "input.txt"
    input_path.write_text("content", encoding="utf-8")

    result = main.run({"input_path": str(input_path)})

    assert result["ok"] is False
    assert result["error"] == "input_path 必须指向 .md 或 .markdown 文件。"
    assert not input_path.with_suffix(".docx").exists()


def test_tool_defaults_to_portrait_and_supports_whole_document_landscape(
    tmp_path: Path,
) -> None:
    input_path = tmp_path / "input.md"
    input_path.write_text("# 页面方向", encoding="utf-8")
    portrait_path = tmp_path / "portrait.docx"
    landscape_path = tmp_path / "landscape.docx"

    portrait_result = main.run(
        {"input_path": str(input_path), "output_path": str(portrait_path)}
    )
    landscape_result = main.run(
        {
            "input_path": str(input_path),
            "output_path": str(landscape_path),
            "page_orientation": "landscape",
        }
    )

    assert portrait_result["ok"] is True
    assert landscape_result["ok"] is True
    portrait = Document(portrait_path).sections[0]
    landscape = Document(landscape_path).sections[0]
    assert portrait.orientation == WD_ORIENT.PORTRAIT
    assert portrait.page_width < portrait.page_height
    assert landscape.orientation == WD_ORIENT.LANDSCAPE
    assert landscape.page_width > landscape.page_height


def test_tool_rejects_unknown_page_orientation_without_creating_output(
    tmp_path: Path,
) -> None:
    input_path = tmp_path / "input.md"
    output_path = tmp_path / "output.docx"
    input_path.write_text("content", encoding="utf-8")

    result = main.run(
        {
            "input_path": str(input_path),
            "output_path": str(output_path),
            "page_orientation": "auto",
        }
    )

    assert result["ok"] is False
    assert result["error"] == "page_orientation 必须为 portrait 或 landscape。"
    assert not output_path.exists()


def test_page_orientation_schema_matches_runtime_contract() -> None:
    schema = json.loads(
        (TOOL_ROOT / ".tool" / "input.schema.json").read_text(encoding="utf-8")
    )
    orientation = schema["properties"]["page_orientation"]

    assert orientation["enum"] == ["portrait", "landscape"]
    assert orientation["default"] == "portrait"


def test_tool_defaults_to_letter_and_supports_explicit_a4(tmp_path: Path) -> None:
    input_path = tmp_path / "input.md"
    input_path.write_text("# 纸张", encoding="utf-8")
    letter_path = tmp_path / "letter.docx"
    a4_path = tmp_path / "a4.docx"

    assert main.run({"input_path": str(input_path), "output_path": str(letter_path)})["ok"]
    assert main.run(
        {
            "input_path": str(input_path),
            "output_path": str(a4_path),
            "page_size": "a4",
        }
    )["ok"]
    letter = Document(letter_path).sections[0]
    a4 = Document(a4_path).sections[0]
    assert letter.page_width.inches == pytest.approx(8.5, abs=0.01)
    assert letter.page_height.inches == pytest.approx(11, abs=0.01)
    assert a4.page_width.mm == pytest.approx(210, abs=0.2)
    assert a4.page_height.mm == pytest.approx(297, abs=0.2)


def test_tool_rejects_unknown_page_size_without_creating_output(tmp_path: Path) -> None:
    input_path = tmp_path / "input.md"
    output_path = tmp_path / "output.docx"
    input_path.write_text("content", encoding="utf-8")

    result = main.run(
        {
            "input_path": str(input_path),
            "output_path": str(output_path),
            "page_size": "legal",
        }
    )

    assert result["ok"] is False
    assert result["error"] == "page_size 必须为 a4 或 letter。"
    assert not output_path.exists()


def test_page_size_and_output_schemas_match_runtime_contract() -> None:
    input_schema = json.loads(
        (TOOL_ROOT / ".tool" / "input.schema.json").read_text(encoding="utf-8")
    )
    output_schema = json.loads(
        (TOOL_ROOT / ".tool" / "output.schema.json").read_text(encoding="utf-8")
    )

    page_size = input_schema["properties"]["page_size"]
    assert page_size["enum"] == ["letter", "a4"]
    assert page_size["default"] == "letter"
    assert len(output_schema["oneOf"]) == 4
    assert {branch["properties"]["ok"]["const"] for branch in output_schema["oneOf"]} == {
        True,
        False,
    }
    success_actions = {
        branch["properties"]["action"]["const"]
        for branch in output_schema["oneOf"]
        if branch["properties"]["ok"]["const"] is True
    }
    assert success_actions == {"convert", "extract_template", "list_templates"}


def test_template_actions_extract_list_and_convert(
    tmp_path: Path,
    monkeypatch,
) -> None:
    source_path = tmp_path / "客户样本.docx"
    markdown_path = tmp_path / "报告.md"
    output_path = tmp_path / "报告.docx"
    source = Document()
    source.sections[0].top_margin = Mm(25.4)
    normal = source.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(
        qn("w:eastAsia"),
        "仿宋",
    )
    normal.font.size = Pt(10.5)
    source.add_paragraph("模板正文")
    source.save(source_path)
    markdown_path.write_text("# 标题\n\n正文。", encoding="utf-8")
    monkeypatch.setattr(main, "TEMPLATES_DIR", tmp_path / "assets" / "templates")

    extracted = main.run(
        {
            "action": "extract_template",
            "template_source_path": str(source_path),
            "template_name": "客户模板",
        }
    )
    assert extracted["ok"] is True
    template_id = extracted["template"]["template_id"]

    listed = main.run({"action": "list_templates"})
    assert [item["name"] for item in listed["templates"]] == [
        "内置默认样式",
        "客户模板",
    ]

    converted = main.run(
        {
            "input_path": str(markdown_path),
            "output_path": str(output_path),
            "template_id": template_id,
        }
    )
    assert converted["ok"] is True
    assert converted["template_name"] == "客户模板"
    assert output_path.is_file()
