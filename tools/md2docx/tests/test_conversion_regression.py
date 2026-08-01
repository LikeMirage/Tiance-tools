from __future__ import annotations

import base64
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from lxml import etree

from converter import convert_markdown_to_docx


ONE_PIXEL_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9Y9Z3jcAAAAASUVORK5CYII="
)


def test_rich_markdown_keeps_core_document_features(tmp_path: Path) -> None:
    image_path = tmp_path / "pixel.png"
    image_path.write_bytes(ONE_PIXEL_PNG)
    document, warnings = _convert(
        tmp_path,
        r"""# 标题

正文包含 **粗体**、*斜体*、[链接](https://example.com)、行内公式 $x+1$ 和脚注[^one]。

- 一级项目
  - 二级项目

| 名称 | 公式 | 对齐 |
| --- | --- | :---: |
| 单元格 | $a\mid b$ | 居中 |

![示例](pixel.png)

[^one]: 脚注内容。
""",
    )

    assert warnings == []
    assert document.paragraphs[0].text == "标题"
    assert len(document.tables) == 1
    assert _table_rows(document.tables[0]) == [
        ["名称", "公式", "对齐"],
        ["单元格", "", "居中"],
    ]
    assert len(document._element.xpath(".//m:oMath")) == 2
    assert len(document._element.xpath(".//w:footnoteReference")) == 1
    assert len(document.inline_shapes) == 1
    list_paragraphs = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text in {"一级项目", "二级项目"}
    ]
    assert [paragraph.text for paragraph in list_paragraphs] == ["一级项目", "二级项目"]
    assert [
        paragraph._p.xpath("./w:pPr/w:numPr/w:ilvl/@w:val")[0]
        for paragraph in list_paragraphs
    ] == ["0", "1"]


def test_visible_double_quotes_are_normalized_without_touching_code_or_urls(
    tmp_path: Path,
) -> None:
    document, warnings = _convert(
        tmp_path,
        r"""中文"内容"，English "content"。

"**跨格式**"；**"粗体"**；["链接"](https://example.com/?q="raw")

行内代码 `config="raw"`，尺寸 15"，显式 \"直引号\"，孤立 "引号。

| 名称 |
| --- |
| "表格内容" |

```json
{"key": "value"}
```
""",
    )

    assert warnings == []
    assert document.paragraphs[0].text == "中文“内容”，English “content”。"
    assert document.paragraphs[1].text == "“跨格式”；“粗体”；“链接”"
    assert document.paragraphs[2].text == (
        '行内代码 config="raw"，尺寸 15"，显式 "直引号"，孤立 "引号。'
    )
    assert document.paragraphs[3].text == '{"key": "value"}'
    assert document.tables[0].cell(1, 0).text == "“表格内容”"
    assert any(
        relationship.target_ref == 'https://example.com/?q="raw"'
        for relationship in document.part.rels.values()
    )


def test_invalid_formula_degrades_without_consuming_document_tail(tmp_path: Path) -> None:
    document, warnings = _convert(
        tmp_path,
        r"""$$\frac{a+b}{c+d$$

尾部哨兵 $x+1$。
""",
    )

    assert len(warnings) == 1
    assert "花括号未配对" in warnings[0]
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert r"\frac{a+b}{c+d" in text
    assert "尾部哨兵" in text
    assert len(document._element.xpath(".//m:oMath")) == 1


def test_unclosed_display_formula_keeps_delimiter_literal(tmp_path: Path) -> None:
    document, warnings = _convert(
        tmp_path,
        "\\[\n\n后续内容仍然存在。",
    )

    assert warnings == [r"公式定界符未闭合，已保留原文：\["]
    assert [paragraph.text for paragraph in document.paragraphs] == [
        r"\[",
        "后续内容仍然存在。",
    ]


def test_unclosed_html_keeps_source_and_following_content(tmp_path: Path) -> None:
    document, warnings = _convert(
        tmp_path,
        """<div data-test=\"open\">
未闭合内容

## 后续标题

最终内容
""",
    )

    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert any("HTML 块未闭合" in warning for warning in warnings)
    assert "未闭合内容" in text
    assert "后续标题" in text
    assert "最终内容" in text


def test_word_table_grid_uses_measured_wrap_optimized_widths(tmp_path: Path) -> None:
    document, warnings = _convert(
        tmp_path,
        """| ID | 说明 | 状态 |
| --- | --- | --- |
| 1 | The central limit theorem describes convergence for independent random variables. | OK |
| 2 | 短说明 | OK |
""",
    )

    assert warnings == []
    table = document.tables[0]
    grid_widths = [
        int(column.get(qn("w:w")))
        for column in table._tbl.xpath("./w:tblGrid/w:gridCol")
    ]
    assert len(grid_widths) == 3
    assert sum(grid_widths) == 10080
    assert grid_widths[1] > grid_widths[0] * 3
    assert grid_widths[1] > grid_widths[2] * 3
    assert table._tbl.tblPr.find(qn("w:tblLayout")).get(qn("w:type")) == "fixed"


def test_landscape_table_uses_the_wider_whole_document_page(tmp_path: Path) -> None:
    output_path = tmp_path / "landscape.docx"
    warnings = convert_markdown_to_docx(
        """| ID | Description | Status |
| --- | --- | --- |
| 1 | A deliberately long description for width verification. | OK |
""",
        output_path,
        base_path=tmp_path,
        page_orientation="landscape",
    )

    assert warnings == []
    document = Document(output_path)
    section = document.sections[0]
    widths = [
        int(column.get(qn("w:w")))
        for column in document.tables[0]._tbl.xpath("./w:tblGrid/w:gridCol")
    ]
    assert section.orientation == WD_ORIENT.LANDSCAPE
    assert section.page_width > section.page_height
    assert sum(widths) == 13680


def test_block_boundaries_and_internal_navigation_keep_native_word_structure(
    tmp_path: Path,
) -> None:
    document, warnings = _convert(
        tmp_path,
        "正文第一行  \n"
        "正文第二行\n"
        "[TOC]\n\n"
        "# 标题\n\n"
        "[跳转](#标题)\n\n"
        "---\n",
    )

    assert warnings == []
    assert len(document.paragraphs[0]._p.xpath(".//w:br")) == 1
    assert len(document._element.xpath(".//w:fldSimple[contains(@w:instr, 'TOC')]")) == 1
    hyperlink = document._element.xpath(".//w:hyperlink[@w:anchor]")[0]
    bookmark = document._element.xpath(".//w:bookmarkStart[@w:name]")[0]
    assert hyperlink.get(qn("w:anchor")) == bookmark.get(qn("w:name"))
    assert hyperlink.get(qn("r:id")) is None
    assert len(document._element.xpath(".//w:pPr/w:pBdr/w:bottom")) == 1
    assert document.settings.element.xpath("./w:updateFields/@w:val") == ["true"]


def test_invalid_xml_characters_are_replaced_once_and_reported(tmp_path: Path) -> None:
    document, warnings = _convert(tmp_path, "前\x00后\x01")

    assert document.paragraphs[0].text == "前�后�"
    assert warnings == ["已替换 2 个不能写入 Word XML 的控制字符。"]


def test_native_ordered_list_keeps_start_value_and_valid_numbering_order(
    tmp_path: Path,
) -> None:
    document, warnings = _convert(tmp_path, "3. 第三项\n4. 第四项")

    assert warnings == []
    paragraphs = [paragraph for paragraph in document.paragraphs if paragraph.text]
    num_ids = [paragraph._p.xpath("./w:pPr/w:numPr/w:numId/@w:val")[0] for paragraph in paragraphs]
    assert len(set(num_ids)) == 1
    with ZipFile(tmp_path / "output.docx", "r") as archive:
        numbering = etree.fromstring(archive.read("word/numbering.xml"))
    children = [etree.QName(child).localname for child in numbering]
    assert children == sorted(children, key=lambda name: {"numPicBullet": 0, "abstractNum": 1, "num": 2, "numIdMacAtCleanup": 3}[name])
    overrides = numbering.xpath(
        ".//w:num[@w:numId=$num_id]/w:lvlOverride/w:startOverride/@w:val",
        namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"},
        num_id=num_ids[0],
    )
    assert overrides == ["3"]


def test_parent_numbering_continues_after_a_nested_list(tmp_path: Path) -> None:
    document, warnings = _convert(
        tmp_path,
        "1. 父项一\n  - 子项\n1. 父项二",
    )

    assert warnings == []
    paragraphs = [paragraph for paragraph in document.paragraphs if paragraph.text]
    num_ids = [paragraph._p.xpath("./w:pPr/w:numPr/w:numId/@w:val")[0] for paragraph in paragraphs]
    assert num_ids[0] == num_ids[2]
    assert num_ids[1] != num_ids[0]


def test_formula_font_size_follows_heading_and_table_context(tmp_path: Path) -> None:
    document, warnings = _convert(
        tmp_path,
        """# 标题 $x+1$

| 公式 |
| --- |
| $y+2$ |
""",
    )

    assert warnings == []
    heading_sizes = document.paragraphs[0]._p.xpath(".//m:oMath//w:rPr/w:sz/@w:val")
    table_sizes = document.tables[0].cell(1, 0)._tc.xpath(".//m:oMath//w:rPr/w:sz/@w:val")
    assert set(heading_sizes) == {"36"}
    assert set(table_sizes) == {"21"}


def _convert(tmp_path: Path, markdown: str) -> tuple[Document, list[str]]:
    output_path = tmp_path / "output.docx"
    warnings = convert_markdown_to_docx(
        markdown,
        output_path,
        base_path=tmp_path,
    )
    return Document(output_path), warnings


def _table_rows(table) -> list[list[str]]:
    return [[cell.text for cell in row.cells] for row in table.rows]
