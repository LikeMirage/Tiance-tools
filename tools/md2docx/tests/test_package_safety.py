from __future__ import annotations

import ctypes
import os
import warnings
from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

import docx_package
import docx_package_xml as package_xml


SE_FILE_OBJECT = 1
DACL_SECURITY_INFORMATION = 0x00000004
SE_DACL_PROTECTED = 0x1000


def test_atomic_save_failure_preserves_existing_file(tmp_path: Path, monkeypatch) -> None:
    output_path = tmp_path / "existing.docx"
    original = b"ORIGINAL"
    output_path.write_bytes(original)

    def fail_postprocess(*args, **kwargs) -> None:
        raise RuntimeError("injected failure")

    monkeypatch.setattr(docx_package, "postprocess_docx_package", fail_postprocess)

    with pytest.raises(RuntimeError, match="injected failure"):
        docx_package.save_document_atomically(
            Document(),
            output_path,
            footnotes=[],
            endnotes=[],
            update_fields=False,
            overwrite=True,
        )

    assert output_path.read_bytes() == original
    assert list(tmp_path.glob(".existing-*")) == []


def test_atomic_no_overwrite_rejects_a_file_created_during_conversion(
    tmp_path: Path,
    monkeypatch,
) -> None:
    output_path = tmp_path / "raced.docx"
    original_postprocess = docx_package.postprocess_docx_package

    def create_competing_output(path, **options) -> None:
        original_postprocess(path, **options)
        output_path.write_bytes(b"OTHER PROCESS")

    monkeypatch.setattr(docx_package, "postprocess_docx_package", create_competing_output)

    with pytest.raises(ValueError, match="目标文件已存在"):
        docx_package.save_document_atomically(
            Document(),
            output_path,
            footnotes=[],
            endnotes=[],
            update_fields=False,
            overwrite=False,
        )

    assert output_path.read_bytes() == b"OTHER PROCESS"
    assert list(tmp_path.glob(".raced-*")) == []


@pytest.mark.skipif(os.name != "nt", reason="Windows ACL regression")
def test_atomic_save_output_inherits_parent_directory_acl(tmp_path: Path) -> None:
    output_path = tmp_path / "output.docx"

    docx_package.save_document_atomically(
        Document(),
        output_path,
        footnotes=[],
        endnotes=[],
        update_fields=False,
        overwrite=False,
    )

    assert not _is_dacl_protected(output_path)


def test_package_validation_rejects_invalid_omml_structure(tmp_path: Path) -> None:
    output_path = tmp_path / "invalid-omml.docx"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph._p.append(
        etree.fromstring(
            b'<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
            b"<m:f><m:num><m:e/></m:num></m:f></m:oMath>"
        )
    )
    document.save(output_path)

    with pytest.raises(ValueError, match="包含无效 Word 公式结构"):
        docx_package.validate_docx_package(output_path)


def test_package_validation_rejects_invalid_word_child_order(tmp_path: Path) -> None:
    output_path = tmp_path / "invalid-order.docx"
    document = Document()
    run = document.add_paragraph().add_run("text")
    properties = run._r.get_or_add_rPr()
    properties.append(OxmlElement("w:sz"))
    properties.append(OxmlElement("w:rFonts"))
    document.save(output_path)

    with pytest.raises(ValueError, match="w:rFonts 不能位于 w:sz 之后"):
        docx_package.validate_docx_package(output_path)


def test_package_validation_rejects_duplicate_zip_entries(tmp_path: Path) -> None:
    output_path = tmp_path / "duplicate.docx"
    Document().save(output_path)
    with ZipFile(output_path, "a") as archive, warnings.catch_warnings():
        warnings.simplefilter("ignore", UserWarning)
        archive.writestr("word/document.xml", archive.read("word/document.xml"))

    with pytest.raises(ValueError, match="重复 ZIP 条目：word/document.xml"):
        docx_package.validate_docx_package(output_path)


def test_note_parts_have_relationships_content_types_styles_and_update_fields(
    tmp_path: Path,
) -> None:
    output_path = tmp_path / "notes.docx"
    docx_package.save_document_atomically(
        Document(),
        output_path,
        footnotes=[(1, "脚注")],
        endnotes=[(1, "尾注")],
        update_fields=True,
        overwrite=False,
    )

    with ZipFile(output_path, "r") as archive:
        names = set(archive.namelist())
        relationships = etree.fromstring(archive.read("word/_rels/document.xml.rels"))
        content_types = etree.fromstring(archive.read("[Content_Types].xml"))
        styles = etree.fromstring(archive.read("word/styles.xml"))
        settings = etree.fromstring(archive.read("word/settings.xml"))
    assert {"word/footnotes.xml", "word/endnotes.xml"} <= names
    relationship_types = relationships.xpath(
        "./pr:Relationship/@Type",
        namespaces={"pr": package_xml.RELATIONSHIPS_NS},
    )
    assert package_xml.FOOTNOTE_RELATIONSHIP in relationship_types
    assert package_xml.ENDNOTE_RELATIONSHIP in relationship_types
    part_names = content_types.xpath(
        "./ct:Override/@PartName",
        namespaces={"ct": package_xml.CONTENT_TYPES_NS},
    )
    assert "/word/footnotes.xml" in part_names
    assert "/word/endnotes.xml" in part_names
    style_ids = styles.xpath(
        "./w:style/@w:styleId",
        namespaces={"w": package_xml.WORD_NS},
    )
    assert {"FootnoteReference", "FootnoteText", "EndnoteReference", "EndnoteText"} <= set(style_ids)
    assert settings.xpath(
        "./w:updateFields/@w:val",
        namespaces={"w": package_xml.WORD_NS},
    ) == ["true"]


def _is_dacl_protected(path: Path) -> bool:
    security_descriptor = ctypes.c_void_p()
    advapi32 = ctypes.windll.advapi32
    advapi32.GetNamedSecurityInfoW.argtypes = [
        ctypes.c_wchar_p,
        ctypes.c_uint,
        ctypes.c_uint,
        ctypes.c_void_p,
        ctypes.c_void_p,
        ctypes.c_void_p,
        ctypes.c_void_p,
        ctypes.POINTER(ctypes.c_void_p),
    ]
    result = advapi32.GetNamedSecurityInfoW(
        str(path),
        SE_FILE_OBJECT,
        DACL_SECURITY_INFORMATION,
        None,
        None,
        None,
        None,
        ctypes.byref(security_descriptor),
    )
    if result:
        raise OSError(result, "无法读取测试文件 ACL")
    try:
        control = ctypes.c_ushort()
        revision = ctypes.c_ulong()
        if not advapi32.GetSecurityDescriptorControl(
            security_descriptor,
            ctypes.byref(control),
            ctypes.byref(revision),
        ):
            raise ctypes.WinError()
        return bool(control.value & SE_DACL_PROTECTED)
    finally:
        ctypes.windll.kernel32.LocalFree(security_descriptor)
